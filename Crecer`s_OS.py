import os
import random
import subprocess
import requests
import configparser
import json
import shutil
import zipfile
import sqlite3
import xml.etree.ElementTree as ET
import tempfile
import winreg
from functools import partial
import matplotlib

matplotlib.use("QtAgg")
from io import BytesIO
from random import choice
from datetime import datetime

import psutil  # Для информации о батарее и сети
from PIL import Image

from PyQt6.QtWidgets import (
    QMainWindow, QTextEdit, QLineEdit,
    QGridLayout, QSpinBox, QCheckBox, QComboBox, QFileDialog,
    QScrollArea, QCalendarWidget, QTimeEdit, QFormLayout,
    QToolButton, QHeaderView, QStyle, QProgressDialog, QAbstractItemView, QSlider
)
from PyQt6.QtWidgets import QLabel, QWidget, QMenu, \
    QGraphicsOpacityEffect
from PyQt6.QtCore import QPoint, QRect, QPropertyAnimation, QEasingCurve, pyqtProperty
from PyQt6.QtGui import QColor, QRadialGradient
from PyQt6.QtGui import QPixmap, QPalette, QBrush, QIcon, QCursor, QMouseEvent, QImage, \
    QPainter, QTransform, QAction
from PyQt6.QtCore import QTimer, QTime, QDate, QSize, QEvent, QPointF, QDir, QFileInfo, QSettings, \
    QStorageInfo
from PyQt6.QtMultimedia import QMediaPlayer, QAudioOutput
from PyQt6.QtMultimediaWidgets import QVideoWidget

import sys
from PyQt6.QtWidgets import (
    QApplication, QDialog, QVBoxLayout, QHBoxLayout, QPushButton, QTabWidget,
    QProgressBar, QMessageBox, QInputDialog, QListWidget, QListWidgetItem,
    QTableWidget, QTableWidgetItem
)
from PyQt6.QtCore import QUrl, Qt
from PyQt6.QtGui import QKeySequence, QShortcut
from PyQt6.QtWebEngineWidgets import QWebEngineView
from PyQt6.QtWebEngineCore import QWebEngineProfile, QWebEnginePage, QWebEngineDownloadRequest, QWebEngineSettings

# PDF поддержка (если доступна)
try:
    from PyQt6.QtPdfWidgets import QPdfView
    from PyQt6.QtPdf import QPdfDocument

    pdf_supported = True
except ImportError:
    pdf_supported = False

try:
    from PyQt6.QtWebEngineWidgets import QWebEngineView

    html_supported = True
except ImportError:
    html_supported = False

SETTINGS_FILE = "setting.json"


def load_json_settings():
    try:
        with open(SETTINGS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


def save_json_settings(settings):
    with open(SETTINGS_FILE, "w", encoding="utf-8") as f:
        json.dump(settings, f, ensure_ascii=False, indent=4)


try:
    import fitz
except ImportError:
    fitz = None
try:
    import docx
except ImportError:
    docx = None
try:
    import openpyxl
except ImportError:
    openpyxl = None
try:
    import pptx
except ImportError:
    pptx = None

try:
    import sounddevice as sd
    import numpy as np
except ImportError:
    sd = None
    np = None


class SnapIndicatorOverlay(QWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.Tool)
        self.setAttribute(Qt.WidgetAttribute.WA_TransparentForMouseEvents)
        self.setAttribute(Qt.WidgetAttribute.WA_NoSystemBackground)
        self.setAttribute(Qt.WidgetAttribute.WA_TranslucentBackground)
        self.dot_radius = 20
        self.dots = []
        self.calculate_dots()

        self.opacity_effect = QGraphicsOpacityEffect(self)
        self.setGraphicsEffect(self.opacity_effect)
        self.opacity_effect.setOpacity(0)

        self._dot_scale = 0.0

    def getDotScale(self):
        return self._dot_scale

    def setDotScale(self, scale):
        self._dot_scale = scale
        self.update()

    dotScale = pyqtProperty(float, fget=getDotScale, fset=setDotScale)

    def calculate_dots(self):
        screen_geom = QApplication.primaryScreen().availableGeometry()
        w = screen_geom.width()
        h = screen_geom.height()
        self.dots = [
            QPoint(0, 0),
            QPoint(w // 2, 0),
            QPoint(w, 0),
            QPoint(0, h // 2),
            QPoint(w, h // 2),
            QPoint(0, h),
            QPoint(w // 2, h),
            QPoint(w, h)
        ]

    def showOverlay(self):
        self.calculate_dots()
        self.setGeometry(QApplication.primaryScreen().availableGeometry())
        self.show()
        self.opacity_anim = QPropertyAnimation(self.opacity_effect, b"opacity")
        self.opacity_anim.setDuration(500)
        self.opacity_anim.setStartValue(0)
        self.opacity_anim.setEndValue(1)
        self.opacity_anim.setEasingCurve(QEasingCurve.Type.InOutQuad)
        self.opacity_anim.start()

        self.scale_anim = QPropertyAnimation(self, b"dotScale")
        self.scale_anim.setDuration(500)
        self.scale_anim.setStartValue(0)
        self.scale_anim.setEndValue(1)
        self.scale_anim.setEasingCurve(QEasingCurve.Type.InOutQuad)
        self.scale_anim.start()

    def hideOverlay(self):
        self.opacity_anim = QPropertyAnimation(self.opacity_effect, b"opacity")
        self.opacity_anim.setDuration(500)
        self.opacity_anim.setStartValue(self.opacity_effect.opacity())
        self.opacity_anim.setEndValue(0)
        self.opacity_anim.setEasingCurve(QEasingCurve.Type.InOutQuad)
        self.opacity_anim.start()

        self.scale_anim = QPropertyAnimation(self, b"dotScale")
        self.scale_anim.setDuration(500)
        self.scale_anim.setStartValue(self._dot_scale)
        self.scale_anim.setEndValue(0)
        self.scale_anim.setEasingCurve(QEasingCurve.Type.InOutQuad)
        self.scale_anim.start()
        self.scale_anim.finished.connect(self.hide)

    def paintEvent(self, event):
        painter = QPainter(self)
        painter.setRenderHint(QPainter.RenderHint.Antialiasing)
        for dot in self.dots:
            gradient = QRadialGradient(QPointF(dot), self.dot_radius * self._dot_scale)
            gradient.setColorAt(0, QColor(0, 120, 215, 220))
            gradient.setColorAt(1, QColor(0, 120, 215, 0))
            painter.setBrush(gradient)
            painter.setPen(Qt.PenStyle.NoPen)
            painter.drawEllipse(QPointF(dot), self.dot_radius * self._dot_scale, self.dot_radius * self._dot_scale)


class DraggableDialog(QDialog):
    SNAP_THRESHOLD = 30
    RESIZE_MARGIN = 10

    def __init__(self, title=""):
        super().__init__()
        self.setWindowTitle(title if title else "Проводник")
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint | Qt.WindowType.Window)
        self._oldPos = None
        self._is_maximized_custom = False
        self.normal_geometry = None

        self.locked = False
        self.always_on_top = False
        self.resizing = False
        self._resize_start_pos = None
        self._resize_start_geom = None

        self.snap_animation = QPropertyAnimation(self, b"geometry")
        self.snap_animation.setDuration(300)
        self.snap_animation.setEasingCurve(QEasingCurve.Type.OutCubic)
        self.resize_animation = QPropertyAnimation(self, b"geometry")
        self.resize_animation.setDuration(400)
        self.resize_animation.setEasingCurve(QEasingCurve.Type.InOutQuad)

        self.main_layout = QVBoxLayout(self)
        self.main_layout.setContentsMargins(0, 0, 0, 0)
        self.main_layout.setSpacing(0)

        self.title_bar = QWidget(self)
        self.title_bar.setStyleSheet("background-color: #444; color: white;")
        title_layout = QHBoxLayout(self.title_bar)
        title_layout.setContentsMargins(5, 2, 5, 2)
        self.title_label = QLabel(title if title else "Проводник")
        title_layout.addWidget(self.title_label)
        title_layout.addStretch()

        self.top_button = QPushButton("⇧")
        self.top_button.setFixedSize(20, 20)
        self.top_button.clicked.connect(self.toggle_always_on_top)
        title_layout.addWidget(self.top_button)

        self.lock_button = QPushButton("🔓")
        self.lock_button.setFixedSize(20, 20)
        self.lock_button.clicked.connect(self.toggle_lock)
        title_layout.addWidget(self.lock_button)

        self.snap_button = QPushButton("☰")
        self.snap_button.setFixedSize(20, 20)
        self.snap_button.clicked.connect(self.open_snap_menu)
        title_layout.addWidget(self.snap_button)

        self.min_button = QPushButton("-")
        self.min_button.setFixedSize(20, 20)
        self.min_button.clicked.connect(self.showMinimized)
        title_layout.addWidget(self.min_button)

        self.max_button = QPushButton("□")
        self.max_button.setFixedSize(20, 20)
        self.max_button.clicked.connect(self.maximize_to_available)
        title_layout.addWidget(self.max_button)

        self.close_button = QPushButton("X")
        self.close_button.setFixedSize(20, 20)
        self.close_button.clicked.connect(self.close)
        title_layout.addWidget(self.close_button)

        self.main_layout.addWidget(self.title_bar)
        self.content_area = QWidget(self)
        self.main_layout.addWidget(self.content_area)

        self.snap_overlay = SnapIndicatorOverlay()
        self.snap_overlay.setGeometry(QApplication.primaryScreen().availableGeometry())
        self.snap_overlay.hide()

    def setContentLayout(self, layout):
        self.content_area.setLayout(layout)
        self.content_layout = layout

    def mousePressEvent(self, event):
        if self.locked:
            return super().mousePressEvent(event)
        if event.button() == Qt.MouseButton.LeftButton:
            if not self.resizing:
                self.snap_overlay.showOverlay()
            if event.position().x() >= self.width() - self.RESIZE_MARGIN and event.position().y() >= self.height() - self.RESIZE_MARGIN:
                self.resizing = True
                self._resize_start_pos = event.globalPosition().toPoint()
                self._resize_start_geom = self.geometry()
            else:
                self._oldPos = event.globalPosition().toPoint()
        return super().mousePressEvent(event)

    def mouseMoveEvent(self, event):
        if self.locked:
            return super().mouseMoveEvent(event)
        if self.resizing:
            delta = event.globalPosition().toPoint() - self._resize_start_pos
            new_width = max(self.minimumWidth(), self._resize_start_geom.width() + delta.x())
            new_height = max(self.minimumHeight(), self._resize_start_geom.height() + delta.y())
            self.setGeometry(self._resize_start_geom.x(), self._resize_start_geom.y(), new_width, new_height)
        elif self._oldPos:
            delta = event.globalPosition().toPoint() - self._oldPos
            self.move(self.x() + delta.x(), self.y() + delta.y())
            self._oldPos = event.globalPosition().toPoint()
        else:
            if event.position().x() >= self.width() - self.RESIZE_MARGIN and event.position().y() >= self.height() - self.RESIZE_MARGIN:
                self.setCursor(Qt.CursorShape.SizeFDiagCursor)
            else:
                self.unsetCursor()
        return super().mouseMoveEvent(event)

    def mouseReleaseEvent(self, event):
        if self.locked:
            return super().mouseReleaseEvent(event)
        if self.resizing:
            self.resizing = False
            self.animate_resize_effect()
        else:
            self.check_and_snap()
        self._oldPos = None
        self.snap_overlay.hideOverlay()
        return super().mouseReleaseEvent(event)

    def check_and_snap(self):
        screen_geom = self.screen().availableGeometry()
        x, y, w, h = self.geometry().getRect()
        threshold = self.SNAP_THRESHOLD

        snap_geom = None
        if x <= threshold and y <= threshold:
            snap_geom = QRect(0, 0, screen_geom.width() // 2, screen_geom.height() // 2)
        elif (x + w) >= (screen_geom.width() - threshold) and y <= threshold:
            snap_geom = QRect(screen_geom.width() // 2, 0, screen_geom.width() // 2, screen_geom.height() // 2)
        elif x <= threshold and (y + h) >= (screen_geom.height() - threshold):
            snap_geom = QRect(0, screen_geom.height() // 2, screen_geom.width() // 2, screen_geom.height() // 2)
        elif (x + w) >= (screen_geom.width() - threshold) and (y + h) >= (screen_geom.height() - threshold):
            snap_geom = QRect(screen_geom.width() // 2, screen_geom.height() // 2, screen_geom.width() // 2,
                              screen_geom.height() // 2)
        elif x <= threshold:
            snap_geom = QRect(0, 0, screen_geom.width() // 2, screen_geom.height())
        elif (x + w) >= (screen_geom.width() - threshold):
            snap_geom = QRect(screen_geom.width() // 2, 0, screen_geom.width() // 2, screen_geom.height())
        elif y <= threshold:
            snap_geom = QRect(0, 0, screen_geom.width(), screen_geom.height() // 2)
        elif (y + h) >= (screen_geom.height() - threshold):
            snap_geom = QRect(0, screen_geom.height() // 2, screen_geom.width(), screen_geom.height() // 2)

        if snap_geom:
            self.animate_to(snap_geom)

    def animate_to(self, target_geom):
        self.snap_animation.stop()
        self.snap_animation.setStartValue(self.geometry())
        self.snap_animation.setEndValue(target_geom)
        self.snap_animation.start()

    def animate_resize_effect(self):
        final_geom = self.geometry()
        scale_factor = 1.05
        scaled_geom = QRect(
            final_geom.center().x() - int(final_geom.width() * scale_factor / 2),
            final_geom.center().y() - int(final_geom.height() * scale_factor / 2),
            int(final_geom.width() * scale_factor),
            int(final_geom.height() * scale_factor)
        )
        self.resize_animation.stop()
        self.resize_animation.setStartValue(scaled_geom)
        self.resize_animation.setEndValue(final_geom)
        self.resize_animation.start()

    def maximize_to_available(self):
        TASKBAR_HEIGHT = 50
        extra_height = 55
        if self.parent():
            parent_geom = self.parent().geometry()
            new_rect = (parent_geom.x(), parent_geom.y(), parent_geom.width(),
                        parent_geom.height() - TASKBAR_HEIGHT + extra_height)
        else:
            screen_geom = self.screen().availableGeometry()
            new_rect = (screen_geom.x(), screen_geom.y(), screen_geom.width(),
                        screen_geom.height() - TASKBAR_HEIGHT + extra_height)
        if self._is_maximized_custom:
            if self.normal_geometry is not None:
                self.animate_to(self.normal_geometry)
            self._is_maximized_custom = False
        else:
            self.normal_geometry = self.geometry()
            self.animate_to(QRect(*new_rect))
            self._is_maximized_custom = True

    def open_snap_menu(self):
        menu = QMenu(self)
        actions = {
            "Верхний левый": lambda: self.animate_to(QRect(0, 0, self.screen().availableGeometry().width() // 2,
                                                           self.screen().availableGeometry().height() // 2)),
            "Верхний правый": lambda: self.animate_to(
                QRect(self.screen().availableGeometry().width() // 2, 0, self.screen().availableGeometry().width() // 2,
                      self.screen().availableGeometry().height() // 2)),
            "Нижний левый": lambda: self.animate_to(QRect(0, self.screen().availableGeometry().height() // 2,
                                                          self.screen().availableGeometry().width() // 2,
                                                          self.screen().availableGeometry().height() // 2)),
            "Нижний правый": lambda: self.animate_to(
                QRect(self.screen().availableGeometry().width() // 2, self.screen().availableGeometry().height() // 2,
                      self.screen().availableGeometry().width() // 2, self.screen().availableGeometry().height() // 2)),
            "Левая половина": lambda: self.animate_to(QRect(0, 0, self.screen().availableGeometry().width() // 2,
                                                            self.screen().availableGeometry().height())),
            "Правая половина": lambda: self.animate_to(
                QRect(self.screen().availableGeometry().width() // 2, 0, self.screen().availableGeometry().width() // 2,
                      self.screen().availableGeometry().height())),
            "Верхняя половина": lambda: self.animate_to(QRect(0, 0, self.screen().availableGeometry().width(),
                                                              self.screen().availableGeometry().height() // 2)),
            "Нижняя половина": lambda: self.animate_to(
                QRect(0, self.screen().availableGeometry().height() // 2, self.screen().availableGeometry().width(),
                      self.screen().availableGeometry().height() // 2))
        }
        for text, func in actions.items():
            action = menu.addAction(text)
            action.triggered.connect(func)
        menu.exec(QCursor.pos())

    def toggle_always_on_top(self):
        self.always_on_top = not self.always_on_top
        flags = self.windowFlags()
        if self.always_on_top:
            self.top_button.setText("⇪")
            flags |= Qt.WindowType.WindowStaysOnTopHint
        else:
            self.top_button.setText("⇧")
            flags &= ~Qt.WindowType.WindowStaysOnTopHint
        self.setWindowFlags(flags)
        self.show()

    def toggle_lock(self):
        self.locked = not self.locked
        if self.locked:
            self.lock_button.setText("🔒")
        else:
            self.lock_button.setText("🔓")

    def keyPressEvent(self, event):
        if event.key() in (Qt.Key.Key_Enter, Qt.Key.Key_Return):
            event.ignore()
            return
        super().keyPressEvent(event)


# Скрытое приложение
class HiddenApp(DraggableDialog):
    def __init__(self):
        title = "Скрытое приложение"
        text = "Это скрытое приложение"
        super().__init__(title)
        self.resize(300, 200)
        layout = QVBoxLayout()
        self.setContentLayout(layout)
        label = QLabel(text)
        label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        layout.addWidget(label)
        self.show()


def load_json(filename):
    if os.path.exists(filename):
        try:
            with open(filename, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}
    return {}


def save_json(filename, data):
    with open(filename, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=4)


class App4(DraggableDialog):
    HISTORY_FILE = "history.json"
    BOOKMARKS_FILE = "bookmarks.json"
    PASSWORDS_FILE = "passwords.json"

    def __init__(self):
        super().__init__("Браузер")
        self.resize(1000, 700)

        self.history_data = load_json(self.HISTORY_FILE)
        self.bookmarks = load_json(self.BOOKMARKS_FILE)
        self.passwords = load_json(self.PASSWORDS_FILE)

        main_layout = QVBoxLayout()
        self.setContentLayout(main_layout)

        toolbar = QHBoxLayout()

        self.btn_back = QPushButton("←")
        self.btn_back.clicked.connect(self.navigate_back)
        self.btn_forward = QPushButton("→")
        self.btn_forward.clicked.connect(self.navigate_forward)
        self.btn_refresh = QPushButton("⟳")
        self.btn_refresh.clicked.connect(self.refresh_page)
        self.btn_home = QPushButton("🏠")
        self.btn_home.clicked.connect(self.go_home)
        toolbar.addWidget(self.btn_back)
        toolbar.addWidget(self.btn_forward)
        toolbar.addWidget(self.btn_refresh)
        toolbar.addWidget(self.btn_home)

        self.address_bar = QLineEdit()
        self.address_bar.returnPressed.connect(self.load_address)
        toolbar.addWidget(self.address_bar)

        btn_new_tab = QPushButton("Новая вкладка")
        btn_new_tab.clicked.connect(lambda: self.add_new_tab("https://www.google.com"))
        toolbar.addWidget(btn_new_tab)

        btn_bookmark = QPushButton("Закладки")
        btn_bookmark.clicked.connect(self.show_bookmarks)
        toolbar.addWidget(btn_bookmark)

        btn_history = QPushButton("История")
        btn_history.clicked.connect(self.show_history)
        toolbar.addWidget(btn_history)

        btn_password = QPushButton("Менеджер паролей")
        btn_password.clicked.connect(self.show_password_manager)
        toolbar.addWidget(btn_password)

        main_layout.addLayout(toolbar)

        self.tab_widget = QTabWidget()
        self.tab_widget.tabBar().setFixedHeight(40)
        self.tab_widget.setTabsClosable(True)
        self.tab_widget.setMovable(True)
        self.tab_widget.tabCloseRequested.connect(self.close_tab)
        self.tab_widget.currentChanged.connect(self.on_tab_change)
        main_layout.addWidget(self.tab_widget)

        self.profile = QWebEngineProfile.defaultProfile()
        self.profile.settings().setAttribute(QWebEngineSettings.WebAttribute.LocalContentCanAccessRemoteUrls, False)
        self.profile.setHttpCacheType(QWebEngineProfile.HttpCacheType.DiskHttpCache)

        self.add_new_tab("https://www.google.com")
        self.show()

    def load_address(self):
        url_text = self.address_bar.text().strip()
        if not url_text.startswith("http"):
            url_text = "https://" + url_text
        current_browser = self.tab_widget.currentWidget()
        if current_browser:
            current_browser.setUrl(QUrl(url_text))

    def open_snap_menu(self):
        super().open_snap_menu()

    def add_new_tab(self, url="https://www.google.com"):
        if not isinstance(url, str):
            url = "https://www.google.com"
        browser = QWebEngineView()
        page = QWebEnginePage(self.profile, browser)
        page.settings().setAttribute(QWebEngineSettings.WebAttribute.JavascriptCanOpenWindows, False)
        browser.setPage(page)
        browser.setUrl(QUrl(url))
        index = self.tab_widget.addTab(browser, "Новая вкладка")
        self.tab_widget.setCurrentIndex(index)

        self.history_data[str(index)] = [url]
        self.save_history()

        browser.urlChanged.connect(lambda qurl, idx=index: self.on_url_changed(idx, qurl))
        browser.titleChanged.connect(lambda title, idx=index: self.tab_widget.setTabText(idx, title))

    def on_url_changed(self, index, qurl):
        url = qurl.toString()
        if self.tab_widget.currentIndex() == index:
            self.address_bar.setText(url)
        key = str(index)
        if key in self.history_data:
            if not self.history_data[key] or self.history_data[key][-1] != url:
                self.history_data[key].append(url)
        else:
            self.history_data[key] = [url]
        self.save_history()

    def close_tab(self, index):
        self.tab_widget.removeTab(index)
        key = str(index)
        if key in self.history_data:
            del self.history_data[key]
            self.save_history()

    def on_tab_change(self, index):
        current_browser = self.tab_widget.widget(index)
        if current_browser:
            self.address_bar.setText(current_browser.url().toString())

    def navigate_back(self):
        current_browser = self.tab_widget.currentWidget()
        if current_browser and current_browser.history().canGoBack():
            current_browser.back()

    def navigate_forward(self):
        current_browser = self.tab_widget.currentWidget()
        if current_browser and current_browser.history().canGoForward():
            current_browser.forward()

    def refresh_page(self):
        current_browser = self.tab_widget.currentWidget()
        if current_browser:
            current_browser.reload()

    def go_home(self):
        current_browser = self.tab_widget.currentWidget()
        if current_browser:
            current_browser.setUrl(QUrl("https://www.google.com"))

    def show_bookmarks(self):
        dlg = QDialog(self)
        dlg.setWindowTitle("Закладки")
        dlg.resize(400, 300)
        layout = QVBoxLayout(dlg)
        list_widget = QListWidget()
        for bm in self.bookmarks.get("items", []):
            item = QListWidgetItem(f"{bm['title']} - {bm['url']}")
            list_widget.addItem(item)
        list_widget.itemDoubleClicked.connect(lambda item: self.open_url_in_new_tab(self.parse_bookmark(item.text())))
        layout.addWidget(list_widget)
        btn_layout = QHBoxLayout()
        btn_delete = QPushButton("Удалить выбранное")
        btn_delete.clicked.connect(lambda: self.delete_selected_bookmark(list_widget))
        btn_clear = QPushButton("Очистить закладки")
        btn_clear.clicked.connect(self.clear_bookmarks)
        btn_layout.addWidget(btn_delete)
        btn_layout.addWidget(btn_clear)
        layout.addLayout(btn_layout)
        btn_add = QPushButton("Добавить текущую страницу")
        btn_add.clicked.connect(self.add_current_bookmark)
        layout.addWidget(btn_add)
        dlg.exec()

    def delete_selected_bookmark(self, list_widget):
        selected_items = list_widget.selectedItems()
        if not selected_items:
            return
        for item in selected_items:
            url = self.parse_bookmark(item.text())
            items = self.bookmarks.get("items", [])
            self.bookmarks["items"] = [bm for bm in items if bm["url"] != url]
            list_widget.takeItem(list_widget.row(item))
        save_json(self.BOOKMARKS_FILE, self.bookmarks)

    def clear_bookmarks(self):
        self.bookmarks = {"items": []}
        save_json(self.BOOKMARKS_FILE, self.bookmarks)
        QMessageBox.information(self, "Закладки", "Закладки очищены.")

    def add_current_bookmark(self):
        current_browser = self.tab_widget.currentWidget()
        if current_browser:
            def add_bookmark(title):
                url = current_browser.url().toString()
                bm = {"title": title, "url": url}
                if "items" not in self.bookmarks:
                    self.bookmarks["items"] = []
                self.bookmarks["items"].append(bm)
                save_json(self.BOOKMARKS_FILE, self.bookmarks)
                QMessageBox.information(self, "Закладки", "Страница добавлена в закладки!")

            current_browser.page().runJavaScript("document.title", add_bookmark)

    def parse_bookmark(self, text):
        parts = text.split(" - ")
        return parts[-1] if len(parts) > 1 else text

    def show_history(self):
        dlg = QDialog(self)
        dlg.setWindowTitle("История")
        dlg.resize(400, 300)
        layout = QVBoxLayout(dlg)
        list_widget = QListWidget()
        current_index = self.tab_widget.currentIndex()
        key = str(current_index)
        history = self.history_data.get(key, [])
        for url in history:
            item = QListWidgetItem(url)
            list_widget.addItem(item)
        list_widget.itemDoubleClicked.connect(lambda item: self.open_url_in_new_tab(item.text()))
        layout.addWidget(list_widget)
        btn_layout = QHBoxLayout()
        btn_delete = QPushButton("Удалить выбранное")
        btn_delete.clicked.connect(lambda: self.delete_selected_history(list_widget, key))
        btn_clear = QPushButton("Очистить историю")
        btn_clear.clicked.connect(lambda: self.clear_history(key, list_widget))
        btn_layout.addWidget(btn_delete)
        btn_layout.addWidget(btn_clear)
        layout.addLayout(btn_layout)
        dlg.exec()

    def delete_selected_history(self, list_widget, key):
        selected_items = list_widget.selectedItems()
        if not selected_items:
            return
        history = self.history_data.get(key, [])
        for item in selected_items:
            text = item.text()
            if text in history:
                history.remove(text)
            list_widget.takeItem(list_widget.row(item))
        self.history_data[key] = history
        self.save_history()

    def clear_history(self, key, list_widget):
        self.history_data[key] = []
        self.save_history()
        list_widget.clear()
        QMessageBox.information(self, "История", "История очищена.")

    def open_url_in_new_tab(self, url):
        self.add_new_tab(url)

    def show_password_manager(self):
        dlg = QDialog(self)
        dlg.setWindowTitle("Менеджер паролей")
        dlg.resize(400, 300)
        layout = QVBoxLayout(dlg)
        table = QTableWidget(0, 3)
        table.setHorizontalHeaderLabels(["Сайт", "Пользователь", "Пароль"])
        layout.addWidget(table)
        for site, creds in self.passwords.items():
            row = table.rowCount()
            table.insertRow(row)
            table.setItem(row, 0, QTableWidgetItem(site))
            table.setItem(row, 1, QTableWidgetItem(creds[0]))
            table.setItem(row, 2, QTableWidgetItem(creds[1]))
        btn_layout = QHBoxLayout()
        btn_delete = QPushButton("Удалить выбранное")
        btn_delete.clicked.connect(lambda: self.delete_selected_password(table))
        btn_clear = QPushButton("Очистить все")
        btn_clear.clicked.connect(lambda: self.clear_passwords(table))
        btn_layout.addWidget(btn_delete)
        btn_layout.addWidget(btn_clear)
        layout.addLayout(btn_layout)
        btn_add = QPushButton("Добавить запись")
        btn_add.clicked.connect(lambda: self.add_password_entry(table))
        layout.addWidget(btn_add)
        dlg.exec()

    def delete_selected_password(self, table):
        selected = table.selectedItems()
        if not selected:
            return
        rows = set()
        for item in selected:
            rows.add(item.row())
        for row in sorted(rows, reverse=True):
            site = table.item(row, 0).text()
            if site in self.passwords:
                del self.passwords[site]
            table.removeRow(row)
        save_json(self.PASSWORDS_FILE, self.passwords)

    def clear_passwords(self, table):
        self.passwords = {}
        save_json(self.PASSWORDS_FILE, self.passwords)
        table.setRowCount(0)
        QMessageBox.information(self, "Пароли", "Все записи удалены.")

    def add_password_entry(self, table):
        site, ok1 = QInputDialog.getText(self, "Новый сайт", "Введите название сайта:")
        if not ok1 or not site:
            return
        username, ok2 = QInputDialog.getText(self, "Пользователь", "Введите имя пользователя:")
        if not ok2:
            return
        password, ok3 = QInputDialog.getText(self, "Пароль", "Введите пароль:")
        if not ok3:
            return
        self.passwords[site] = [username, password]
        save_json(self.PASSWORDS_FILE, self.passwords)
        row = table.rowCount()
        table.insertRow(row)
        table.setItem(row, 0, QTableWidgetItem(site))
        table.setItem(row, 1, QTableWidgetItem(username))
        table.setItem(row, 2, QTableWidgetItem(password))

    def save_history(self):
        save_json(self.HISTORY_FILE, self.history_data)


class App5(DraggableDialog):
    def __init__(self):
        super().__init__("Калькулятор")
        self.resize(400, 500)
        self.tabs = QTabWidget()
        self.basic_tab = QWidget()
        self.setup_basic_tab()
        self.tabs.addTab(self.basic_tab, "Обычный")
        self.scientific_tab = QWidget()
        self.setup_scientific_tab()
        self.tabs.addTab(self.scientific_tab, "Инженерный")
        self.unit_tab = QWidget()
        self.setup_unit_tab()
        self.tabs.addTab(self.unit_tab, "Перевод единиц")
        main_layout = QVBoxLayout()
        main_layout.addWidget(self.tabs)
        self.setContentLayout(main_layout)
        self.show()

    def setup_basic_tab(self):
        layout = QVBoxLayout()
        self.basic_display = QLineEdit()
        self.basic_display.setReadOnly(True)
        self.basic_display.setAlignment(Qt.AlignmentFlag.AlignRight)
        layout.addWidget(self.basic_display)
        grid = QGridLayout()
        buttons = [
            ('7', 0, 0), ('8', 0, 1), ('9', 0, 2), ('/', 0, 3),
            ('4', 1, 0), ('5', 1, 1), ('6', 1, 2), ('*', 1, 3),
            ('1', 2, 0), ('2', 2, 1), ('3', 2, 2), ('-', 2, 3),
            ('0', 3, 0), ('.', 3, 1), ('=', 3, 2), ('+', 3, 3),
        ]
        for (text, row, col) in buttons:
            btn = QPushButton(text)
            btn.clicked.connect(lambda checked, t=text: self.basic_button_clicked(t))
            grid.addWidget(btn, row, col)
        layout.addLayout(grid)
        self.basic_tab.setLayout(layout)

    def open_snap_menu(self):
        super().open_snap_menu()

    def basic_button_clicked(self, btn_text):
        if btn_text == '=':
            try:
                result = str(eval(self.basic_display.text()))
                self.basic_display.setText(result)
            except Exception:
                self.basic_display.setText("Ошибка")
        else:
            current = self.basic_display.text()
            self.basic_display.setText(current + btn_text)

    def setup_scientific_tab(self):
        layout = QVBoxLayout()
        self.scientific_display = QLineEdit()
        self.scientific_display.setReadOnly(True)
        self.scientific_display.setAlignment(Qt.AlignmentFlag.AlignRight)
        layout.addWidget(self.scientific_display)
        grid = QGridLayout()
        buttons = [
            ('sin', 0, 0), ('cos', 0, 1), ('tan', 0, 2), ('log', 0, 3),
            ('(', 1, 0), (')', 1, 1), ('^', 1, 2), ('sqrt', 1, 3),
            ('7', 2, 0), ('8', 2, 1), ('9', 2, 2), ('/', 2, 3),
            ('4', 3, 0), ('5', 3, 1), ('6', 3, 2), ('*', 3, 3),
            ('1', 4, 0), ('2', 4, 1), ('3', 4, 2), ('-', 4, 3),
            ('0', 5, 0), ('.', 5, 1), ('=', 5, 2), ('+', 5, 3),
        ]
        for (text, row, col) in buttons:
            btn = QPushButton(text)
            btn.clicked.connect(lambda checked, t=text: self.scientific_button_clicked(t))
            grid.addWidget(btn, row, col)
        layout.addLayout(grid)
        self.scientific_tab.setLayout(layout)

    def scientific_button_clicked(self, btn_text):
        if btn_text == '=':
            try:
                import math
                expr = self.scientific_display.text()
                expr = expr.replace('^', '**')
                expr = expr.replace('sqrt', 'math.sqrt')
                expr = expr.replace('sin', 'math.sin')
                expr = expr.replace('cos', 'math.cos')
                expr = expr.replace('tan', 'math.tan')
                expr = expr.replace('log', 'math.log10')
                result = str(eval(expr))
                self.scientific_display.setText(result)
            except Exception:
                self.scientific_display.setText("Ошибка")
        else:
            current = self.scientific_display.text()
            self.scientific_display.setText(current + btn_text)

    def setup_unit_tab(self):
        layout = QVBoxLayout()
        self.unit_input = QLineEdit()
        self.unit_input.setPlaceholderText("Введите значение")
        layout.addWidget(self.unit_input)
        h_layout = QHBoxLayout()
        self.unit_from = QComboBox()
        self.unit_to = QComboBox()
        units = ["m", "km", "mi", "ft"]
        self.unit_from.addItems(units)
        self.unit_to.addItems(units)
        h_layout.addWidget(self.unit_from)
        h_layout.addWidget(self.unit_to)
        layout.addLayout(h_layout)
        self.convert_button = QPushButton("Конвертировать")
        self.convert_button.clicked.connect(self.convert_units)
        layout.addWidget(self.convert_button)
        self.unit_result = QLineEdit()
        self.unit_result.setReadOnly(True)
        layout.addWidget(self.unit_result)
        self.unit_tab.setLayout(layout)

    def convert_units(self):
        try:
            value = float(self.unit_input.text())
            from_unit = self.unit_from.currentText()
            to_unit = self.unit_to.currentText()
            factors = {"m": 1, "km": 1000, "mi": 1609.34, "ft": 0.3048}
            value_in_m = value * factors[from_unit]
            result = value_in_m / factors[to_unit]
            self.unit_result.setText(str(result))
        except Exception:
            self.unit_result.setText("Ошибка")


class WallpaperButton(QPushButton):
    def mousePressEvent(self, event):
        if event.button() == Qt.MouseButton.LeftButton:
            self.clicked.emit()
        elif event.button() == Qt.MouseButton.RightButton:
            if hasattr(self, "onRightClick") and callable(self.onRightClick):
                self.onRightClick()
        else:
            super().mousePressEvent(event)


class WallpaperChooserDialog(QDialog):
    def __init__(self, wallpapers, parent=None):
        super().__init__(parent, flags=Qt.WindowType.Popup)
        self.setWindowTitle("Выберите обои")
        self.setFixedSize(680, 480)
        layout = QVBoxLayout(self)
        self.list_widget = QListWidget()
        self.list_widget.setViewMode(QListWidget.ViewMode.IconMode)
        self.list_widget.setIconSize(QSize(150, 100))
        self.list_widget.setGridSize(QSize(170, 120))
        self.list_widget.setWrapping(True)
        self.list_widget.setResizeMode(QListWidget.ResizeMode.Adjust)
        self.list_widget.setSpacing(10)
        layout.addWidget(self.list_widget)
        self.delete_button = QPushButton("Удалить выбранные обои")
        self.delete_button.setStyleSheet("background-color: #333; color: white; border-radius: 10px;")
        self.delete_button.clicked.connect(self.delete_selected)
        layout.addWidget(self.delete_button)
        self.wallpapers = wallpapers
        self.populate_list()
        self.list_widget.itemDoubleClicked.connect(self.item_double_clicked)
        self.selected_wallpaper = None

    def populate_list(self):
        self.list_widget.clear()
        for wp in self.wallpapers:
            if os.path.exists(wp):
                pixmap = QPixmap(wp).scaled(150, 100, Qt.AspectRatioMode.KeepAspectRatio,
                                            Qt.TransformationMode.SmoothTransformation)
                item = QListWidgetItem()
                item.setIcon(QIcon(pixmap))
                item.setData(Qt.ItemDataRole.UserRole, wp)
                self.list_widget.addItem(item)

    def item_double_clicked(self, item):
        self.selected_wallpaper = item.data(Qt.ItemDataRole.UserRole)
        self.accept()

    def delete_selected(self):
        selected_items = self.list_widget.selectedItems()
        if not selected_items:
            return
        for item in selected_items:
            try:
                wp = item.data(Qt.ItemDataRole.UserRole)
            except RuntimeError:
                continue
            try:
                os.remove(wp)
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось удалить обои {wp}:\n{str(e)}")
        self.populate_list()
        for item in selected_items:
            try:
                wp = item.data(Qt.ItemDataRole.UserRole)
            except RuntimeError:
                continue
            if wp == self.parent().current_wallpaper:
                self.parent().change_wallpaper()


class WallpaperLinkDialog(QDialog):
    def __init__(self, wallpaper_dir, desktop_window=None):
        super().__init__()
        self.wallpaper_dir = wallpaper_dir
        self.desktop_window = desktop_window
        self.setWindowTitle("Добавить обои по ссылке")
        self.setFixedSize(400, 150)

        self.layout = QVBoxLayout(self)

        self.input_layout = QHBoxLayout()
        self.url_line = QLineEdit()
        self.url_line.setPlaceholderText("Вставьте ссылку на обои...")
        self.check_button = QPushButton("✔")
        self.check_button.setFixedSize(30, 30)
        self.check_button.setStyleSheet("border-radius: 10px;")
        self.check_button.clicked.connect(self.download_wallpaper)
        self.input_layout.addWidget(self.url_line)
        self.input_layout.addWidget(self.check_button)
        self.layout.addLayout(self.input_layout)

        self.status_label = QLabel("")
        self.layout.addWidget(self.status_label)

        self.done_button = QPushButton("Готово")
        self.done_button.setStyleSheet("border-radius: 10px;")
        self.done_button.clicked.connect(self.accept)
        self.layout.addWidget(self.done_button)

    def download_wallpaper(self):
        url = self.url_line.text().strip()
        if not url:
            self.status_label.setText("Введите ссылку.")
            return
        self.status_label.setText("Загрузка...")
        try:
            response = requests.get(url, timeout=10)
            response.raise_for_status()
            image = Image.open(BytesIO(response.content))
            from urllib.parse import urlparse, unquote
            parsed = urlparse(url)
            path = unquote(parsed.path)
            ext = os.path.splitext(path)[1]
            if not ext:
                ext = ".jpg"
            filename = f"link_wallpaper_{datetime.now().strftime('%Y%m%d%H%M%S')}{ext}"
            dest_path = os.path.join(self.wallpaper_dir, filename)
            image.save(dest_path)
            self.status_label.setText("Загрузка успешна")
            if self.desktop_window:
                self.desktop_window.wallpapers.append(dest_path)
        except Exception as e:
            self.status_label.setText("Не удалось загрузить")


class CursorChooserApp(DraggableDialog):
    def __init__(self, parent=None):
        title = "Выбор курсора"
        super().__init__(title)
        self.resize(600, 500)

        self.tabs = QTabWidget()
        self.standard_tab = QWidget()
        self.custom_tab = QWidget()

        standard_layout = QVBoxLayout()
        self.standard_list = QListWidget()
        self.standard_list.setViewMode(QListWidget.ViewMode.IconMode)
        self.standard_list.setIconSize(QSize(32, 32))
        self.standard_list.setGridSize(QSize(50, 50))
        self.standard_list.setSpacing(10)
        standard_layout.addWidget(self.standard_list)
        self.standard_tab.setLayout(standard_layout)

        default_item = QListWidgetItem("По умолчанию")
        default_item.setData(Qt.ItemDataRole.UserRole, "default")
        self.standard_list.addItem(default_item)

        custom_layout = QVBoxLayout()
        self.custom_list = QListWidget()
        self.custom_list.setViewMode(QListWidget.ViewMode.IconMode)
        self.custom_list.setIconSize(QSize(32, 32))
        self.custom_list.setGridSize(QSize(50, 50))
        self.custom_list.setSpacing(10)
        custom_layout.addWidget(self.custom_list)
        btn_layout = QHBoxLayout()
        self.upload_button = QPushButton("Загрузить с компьютера")
        self.upload_button.setStyleSheet("border-radius: 10px;")
        self.upload_button.clicked.connect(self.upload_cursor)
        btn_layout.addWidget(self.upload_button)
        custom_layout.addLayout(btn_layout)
        self.custom_tab.setLayout(custom_layout)

        self.tabs.addTab(self.standard_tab, "Стандартные")
        self.tabs.addTab(self.custom_tab, "Пользовательские")

        main_layout = QVBoxLayout()
        main_layout.addWidget(self.tabs)
        self.setContentLayout(main_layout)

        self.standard_list.itemDoubleClicked.connect(self.standard_item_double_clicked)
        self.custom_list.itemDoubleClicked.connect(self.custom_item_double_clicked)

        self.populate_custom_list()

    def populate_custom_list(self):
        self.custom_list.clear()
        cursor_folder = "cursors"
        if not os.path.exists(cursor_folder):
            os.makedirs(cursor_folder)
        for file in os.listdir(cursor_folder):
            if file.lower().endswith((".png", ".ico", ".cur")):
                path = os.path.join(cursor_folder, file)
                icon = QIcon(path)
                item = QListWidgetItem(icon, file)
                item.setData(Qt.ItemDataRole.UserRole, path)
                self.custom_list.addItem(item)

    def standard_item_double_clicked(self, item):
        data = item.data(Qt.ItemDataRole.UserRole)
        if data == "default":
            QApplication.restoreOverrideCursor()
            if self.parent() is not None:
                self.parent().cursor_path = ""
                self.parent().save_current_cursor()
            self.accept()

    def custom_item_double_clicked(self, item):
        cursor_path = item.data(Qt.ItemDataRole.UserRole)
        QApplication.setOverrideCursor(QCursor(QPixmap(cursor_path)))
        if self.parent() is not None:
            self.parent().cursor_path = cursor_path
            self.parent().save_current_cursor()
        self.accept()

    def upload_cursor(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Выберите файл курсора", "",
                                                   "Images (*.png *.ico *.cur *.jpg *.jpeg *.bmp)")
        if file_path:
            reply = QMessageBox.question(self, "Обрезка изображения", "Применить обрезку изображения до квадрата?",
                                         QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No)
            try:
                image = Image.open(file_path)
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось открыть изображение: {str(e)}")
                return
            if reply == QMessageBox.StandardButton.Yes:
                width, height = image.size
                min_dim = min(width, height)
                left = (width - min_dim) // 2
                top = (height - min_dim) // 2
                right = left + min_dim
                bottom = top + min_dim
                image = image.crop((left, top, right, bottom))
            image = image.resize((32, 32), Image.Resampling.LANCZOS)
            cursor_folder = "cursors"
            if not os.path.exists(cursor_folder):
                os.makedirs(cursor_folder)
            filename = f"cursor_{datetime.now().strftime('%Y%m%d%H%M%S')}" + os.path.splitext(file_path)[1]
            dest_path = os.path.join(cursor_folder, filename)
            try:
                image.save(dest_path)
                QMessageBox.information(self, "Успех", "Курсор успешно загружен.")
                self.populate_custom_list()
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить курсор: {str(e)}")


class CalendarWindow(QDialog):
    NOTES_FILE = "calendar_notes.json"

    def __init__(self, parent=None):
        super().__init__(parent, flags=Qt.WindowType.Popup)
        self.setWindowTitle("Календарь")
        self.setFixedSize(400, 400)
        self.notes = {}
        self.load_notes()
        layout = QVBoxLayout(self)
        self.calendar = QCalendarWidget()
        self.calendar.setGridVisible(True)
        self.calendar.setSelectedDate(QDate.currentDate())
        layout.addWidget(self.calendar)
        self.notes_list = QListWidget()
        layout.addWidget(self.notes_list)
        self.note_edit = QLineEdit()
        self.note_edit.setPlaceholderText("Введите заметку для выбранного времени...")
        layout.addWidget(self.note_edit)
        self.time_edit = QTimeEdit()
        self.time_edit.setDisplayFormat("HH:mm")
        layout.addWidget(self.time_edit)
        self.save_button = QPushButton("Сохранить заметку")
        self.save_button.setStyleSheet("border-radius: 10px;")
        layout.addWidget(self.save_button)
        self.save_button.clicked.connect(self.save_note)
        self.calendar.selectionChanged.connect(self.update_notes)
        self.update_notes()

    def update_notes(self):
        date_str = self.calendar.selectedDate().toString("yyyy-MM-dd")
        self.notes_list.clear()
        if date_str in self.notes:
            for note in self.notes[date_str]:
                self.notes_list.addItem(note)

    def save_note(self):
        date_str = self.calendar.selectedDate().toString("yyyy-MM-dd")
        time_str = self.time_edit.time().toString("HH:mm")
        note_text = self.note_edit.text().strip()
        if note_text:
            full_note = f"{time_str} - {note_text}"
            if date_str not in self.notes:
                self.notes[date_str] = []
            self.notes[date_str].append(full_note)
            self.note_edit.clear()
            self.update_notes()
            self.save_notes()

    def load_notes(self):
        try:
            with open(self.NOTES_FILE, "r", encoding="utf-8") as f:
                self.notes = json.load(f)
        except Exception:
            self.notes = {}

    def save_notes(self):
        with open(self.NOTES_FILE, "w", encoding="utf-8") as f:
            json.dump(self.notes, f, ensure_ascii=False, indent=4)


class TerminalApp(DraggableDialog):
    def __init__(self, desktop_window=None):
        super().__init__("Терминал")
        self.desktop_window = desktop_window
        self.resize(700, 500)
        layout = QVBoxLayout()
        self.setContentLayout(layout)

        self.output = QTextEdit()
        self.output.setReadOnly(True)
        self.output.setStyleSheet("background-color: #1e1e1e; color: #dcdcdc; font-family: Consolas, monospace;")
        layout.addWidget(self.output)

        self.suggestion_list = QListWidget()
        self.suggestion_list.setStyleSheet(
            "background-color: #2e2e2e; color: #dcdcdc; font-family: Consolas, monospace;")
        self.suggestion_list.itemDoubleClicked.connect(self.suggestion_clicked)
        self.suggestion_list.setFixedHeight(0)
        layout.addWidget(self.suggestion_list)

        self.input = QLineEdit()
        self.input.setStyleSheet("background-color: #333; color: #dcdcdc; font-family: Consolas, monospace;")
        self.input.textChanged.connect(self.update_suggestions)
        self.input.returnPressed.connect(self.process_command)
        layout.addWidget(self.input)

        self.commands = [
            "help", "echo", "clear", "exit", "list_apps",
            "open terminal", "open file_explorer", "open settings",
            "wallpaper", "shutdown", "time", "date"
        ]
        self.prank_mode = False
        self._inverting_mouse = False
        self.installEventFilter(self)
        self.input.installEventFilter(self)

        self.print_welcome()

    def print_welcome(self):
        self.output.append("Добро пожаловать в Терминал!")
        self.output.append("Введите 'help' для списка команд.\n")

    def open_snap_menu(self):
        super().open_snap_menu()

    def update_suggestions(self, text):
        self.suggestion_list.clear()
        if text == "":
            self.suggestion_list.setFixedHeight(0)
            return
        matches = [cmd for cmd in self.commands if cmd.startswith(text.lower())]
        for cmd in matches:
            self.suggestion_list.addItem(cmd)
        self.suggestion_list.setFixedHeight(len(matches) * 20 if matches else 0)

    def suggestion_clicked(self, item):
        self.input.setText(item.text())

    def process_command(self):
        cmd = self.input.text().strip()
        if cmd == "open secret":
            if self.desktop_window is not None:
                self.desktop_window.launch_app(self.desktop_window.open_hidden_app, "hidden")
            else:
                self.output.append("Управление приложениями недоступно.")
            self.input.clear()
            self.suggestion_list.clear()
            self.suggestion_list.setFixedHeight(0)
            return

        self.output.append(f"> {cmd}")
        if cmd == "":
            self.input.clear()
            return
        if cmd == "help":
            self.print_help()
        elif cmd.startswith("echo "):
            self.output.append(cmd[5:])
        elif cmd == "clear":
            self.output.clear()
        elif cmd == "exit":
            self.close()
        elif cmd == "list_apps":
            self.output.append("Доступные приложения: terminal, file_explorer, settings, wallpaper, shutdown")
        elif cmd.startswith("open "):
            app_name = cmd[5:].strip().lower()
            if self.desktop_window is not None:
                if app_name == "terminal":
                    self.desktop_window.launch_app(self.desktop_window.open_terminal_app, "terminal")
                elif app_name == "file_explorer":
                    self.desktop_window.launch_app(self.desktop_window.open_file_explorer, "file_explorer")
                elif app_name == "settings":
                    self.desktop_window.launch_app(lambda: self.desktop_window.open_settings(), "settings")
                else:
                    self.output.append("Неизвестное приложение.")
            else:
                self.output.append("Управление приложениями недоступно.")
        elif cmd == "wallpaper":
            if self.desktop_window is not None:
                self.desktop_window.change_wallpaper()
                self.output.append("Обои изменены.")
            else:
                self.output.append("Невозможно изменить обои.")
        elif cmd == "shutdown":
            if self.desktop_window is not None:
                self.output.append("Завершение работы рабочего стола...")
                self.desktop_window.close()
            else:
                self.output.append("Завершение работы недоступно.")
        elif cmd == "time":
            self.output.append("Текущее время: " + QTime.currentTime().toString("HH:mm:ss"))
        elif cmd == "date":
            self.output.append("Текущая дата: " + datetime.now().strftime("%Y-%m-%d"))

        elif cmd == "prank":
            self.prank_mode = not self.prank_mode
            if self.prank_mode:
                self.output.append("Пранк-режим включён!")
            else:
                self.output.append("Пранк-режим отключён!")
        else:
            self.output.append("Неизвестная команда. Введите 'help' для списка команд.")
        self.input.clear()
        self.suggestion_list.clear()
        self.suggestion_list.setFixedHeight(0)

    def print_help(self):
        help_text = (
            "Список команд:\n"
            "help - показать справку\n"
            "echo [текст] - вывести текст\n"
            "clear - очистить экран\n"
            "exit - выйти\n"
            "list_apps - показать список приложений\n"
            "open [приложение] - открыть приложение (terminal, file_explorer, settings)\n"
            "wallpaper - сменить обои\n"
            "shutdown - завершить работу рабочего стола\n"
            "time - показать текущее время\n"
            "date - показать текущую дату\n"
            "prank - переключить пранк-режим\n"
        )
        self.output.append(help_text)

    def eventFilter(self, source, event):
        if self.prank_mode:
            if event.type() == QEvent.Type.KeyPress and source == self.input:
                if event.text() and event.text().isalpha():
                    letters = "абвгдеёжзийклмнопрстуфхцчшщъыьэюя"
                    random_letter = random.choice(letters).upper() if event.text().isupper() else random.choice(letters)
                    self.input.insert(random_letter)
                    return True
            if event.type() in (QEvent.Type.MouseButtonPress, QEvent.Type.MouseButtonRelease, QEvent.Type.MouseMove):
                if not self._inverting_mouse and hasattr(source, "width"):
                    self._inverting_mouse = True
                    pos = event.position().toPoint()
                    new_x = source.width() - pos.x()
                    new_pos = QPointF(new_x, pos.y())
                    new_event = QMouseEvent(
                        event.type(),
                        new_pos,
                        new_pos,
                        event.globalPosition(),
                        event.button(),
                        event.buttons(),
                        event.modifiers()
                    )
                    result = source.event(new_event)
                    self._inverting_mouse = False
                    return result
        return super().eventFilter(source, event)


class SettingsWindow(DraggableDialog):
    CONFIG_FILE = "setting.json"

    def __init__(self, desktop_window, developer=False):
        super().__init__("Настройки")
        self.desktop_window = desktop_window
        self.resize(450, 400)
        self.tabs = QTabWidget()

        # Вкладка "Общие"
        general_tab = QWidget()
        form_general = QFormLayout(general_tab)
        self.auto_wallpaper_checkbox = QCheckBox("Автоматическая смена обоев")
        self.auto_wallpaper_checkbox.setChecked(desktop_window.auto_wallpaper)
        form_general.addRow("", self.auto_wallpaper_checkbox)
        self.interval_spinbox = QSpinBox()
        self.interval_spinbox.setRange(10, 3600)
        self.interval_spinbox.setValue(desktop_window.wallpaper_interval // 1000)
        form_general.addRow("Интервал (сек):", self.interval_spinbox)

        self.theme_combo = QComboBox()
        self.theme_combo.addItems(["dark", "light"])
        self.theme_combo.setCurrentIndex(0)
        form_general.addRow("Тема Crecer's OS:", self.theme_combo)

        self.brightness_slider = QSlider(Qt.Orientation.Horizontal)
        self.brightness_slider.setRange(0, 100)
        self.brightness_slider.setValue(getattr(desktop_window, "brightness", 50))
        form_general.addRow("Яркость экрана:", self.brightness_slider)

        self.tabs.addTab(general_tab, "Общие")

        wallpaper_tab = QWidget()
        form_wallpaper = QVBoxLayout(wallpaper_tab)
        self.add_wallpaper_button = QPushButton("Добавить обои")
        self.add_wallpaper_button.setStyleSheet(
            "background-color: #333; color: white; border: 1px solid #555; border-radius: 10px; padding: 4px 6px; font-size: 16px;")
        self.add_wallpaper_button.clicked.connect(lambda: desktop_window.add_wallpaper())
        form_wallpaper.addWidget(self.add_wallpaper_button)
        self.add_wallpaper_by_link_button = QPushButton("Добавить по ссылке")
        self.add_wallpaper_by_link_button.setStyleSheet(
            "background-color: #333; color: white; border: 1px solid #555; border-radius: 10px; padding: 4px 6px; font-size: 16px;")
        self.add_wallpaper_by_link_button.clicked.connect(self.add_wallpaper_by_link)
        form_wallpaper.addWidget(self.add_wallpaper_by_link_button)
        self.reset_button = QPushButton("Сбросить настройки")
        self.reset_button.setStyleSheet(
            "background-color: #333; color: white; border: 1px solid #555; border-radius: 10px; padding: 4px 6px; font-size: 16px;")
        self.reset_button.clicked.connect(self.reset_settings)
        form_wallpaper.addWidget(self.reset_button)
        self.tabs.addTab(wallpaper_tab, "Обои")

        self.developer_tab = None

        button_layout = QHBoxLayout()
        self.save_button = QPushButton("Сохранить")
        self.save_button.setStyleSheet(
            "background-color: #333; color: white; border: 1px solid #555; border-radius: 10px; padding: 4px 6px; font-size: 16px;")
        self.save_button.clicked.connect(self.save_settings)
        button_layout.addWidget(self.save_button)
        self.cancel_button = QPushButton("Отмена")
        self.cancel_button.setStyleSheet(
            "background-color: #333; color: white; border: 1px solid #555; border-radius: 10px; padding: 4px 6px; font-size: 16px;")
        self.cancel_button.clicked.connect(self.close)
        button_layout.addWidget(self.cancel_button)
        main_layout = QVBoxLayout()
        main_layout.addWidget(self.tabs)
        main_layout.addLayout(button_layout)
        self.setContentLayout(main_layout)
        self.load_settings()

    def maximize_to_available(self):
        super().maximize_to_available()

    def open_snap_menu(self):
        super().open_snap_menu()

    def add_wallpaper_by_link(self):
        dialog = WallpaperLinkDialog(self.desktop_window.wallpaper_dir, desktop_window=self.desktop_window)
        dialog.exec()

    def add_wallpaper(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Выберите обои", "",
                                                   "Images (*.png *.jpg *.jpeg *.bmp *.gif)")
        if file_path:
            base_name = os.path.basename(file_path)
            dest_path = os.path.join(
                self.desktop_window.wallpaper_dir,
                f"custom_{datetime.now().strftime('%Y%m%d%H%M%S')}_{base_name}"
            )
            try:
                with open(file_path, "rb") as src, open(dest_path, "wb") as dst:
                    dst.write(src.read())
                QMessageBox.information(self, "Обои", "Обои успешно добавлены!")
                self.desktop_window.wallpapers.append(dest_path)
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось добавить обои: {str(e)}")

    def load_settings(self):
        settings = load_json_settings()
        self.auto_wallpaper_checkbox.setChecked(settings.get("auto_wallpaper", False))
        self.interval_spinbox.setValue(settings.get("wallpaper_interval", 60))
        theme = settings.get("theme", "dark")
        index = self.theme_combo.findText(theme)
        if index != -1:
            self.theme_combo.setCurrentIndex(index)
        brightness = settings.get("brightness", 50)
        self.brightness_slider.setValue(brightness)
        setattr(self.desktop_window, "brightness", brightness)

    def save_settings(self):
        self.desktop_window.auto_wallpaper = self.auto_wallpaper_checkbox.isChecked()
        self.desktop_window.wallpaper_interval = self.interval_spinbox.value() * 1000
        chosen_theme = self.theme_combo.currentText()
        self.apply_windows_theme(chosen_theme)
        brightness_value = self.brightness_slider.value()
        setattr(self.desktop_window, "brightness", brightness_value)
        if hasattr(self.desktop_window, "apply_brightness"):
            self.desktop_window.apply_brightness(brightness_value)
        settings = {
            "auto_wallpaper": self.desktop_window.auto_wallpaper,
            "wallpaper_interval": self.interval_spinbox.value(),
            "theme": chosen_theme,
            "brightness": brightness_value,
            "current_wallpaper": self.desktop_window.current_wallpaper if getattr(self.desktop_window,
                                                                                  "current_wallpaper", "") else "",
            "cursor": getattr(self.desktop_window, "cursor_path", "")
        }
        save_json_settings(settings)
        if self.desktop_window.auto_wallpaper:
            self.desktop_window.start_wallpaper_timer()
        else:
            self.desktop_window.stop_wallpaper_timer()
        self.close()

    def reset_settings(self):
        self.auto_wallpaper_checkbox.setChecked(False)
        self.interval_spinbox.setValue(60)
        self.theme_combo.setCurrentIndex(0)
        self.brightness_slider.setValue(50)

    def apply_windows_theme(self, theme: str):
        key_path = r"Software\Microsoft\Windows\CurrentVersion\Themes\Personalize"
        try:
            registry_key = winreg.OpenKey(winreg.HKEY_CURRENT_USER, key_path, 0, winreg.KEY_WRITE)
            if theme.lower() == "dark":
                value = 0
                winreg.SetValueEx(registry_key, "AppsUseLightTheme", 0, winreg.REG_DWORD, value)
                winreg.SetValueEx(registry_key, "SystemUsesLightTheme", 0, winreg.REG_DWORD, value)
            elif theme.lower() == "light":
                value = 1
                winreg.SetValueEx(registry_key, "AppsUseLightTheme", 0, winreg.REG_DWORD, value)
                winreg.SetValueEx(registry_key, "SystemUsesLightTheme", 0, winreg.REG_DWORD, value)
            winreg.CloseKey(registry_key)
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось изменить тему: {str(e)}")


class DesktopWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.desktop_title = "Рабочий стол"
        self.start_text = "Пуск"
        self.change_wallpaper_text = "Сменить обои"
        self.file_explorer_text = "Проводник"
        self.terminal_text = "Терминал"
        self.settings_text = "Настройки"
        self.cursor_text = "Курсор"
        self.battery_text = "Батарея"
        self.network_text = "Сеть"

        self.setWindowTitle(self.desktop_title)
        self.setWindowFlags(Qt.WindowType.FramelessWindowHint)
        self.showFullScreen()
        self.auto_wallpaper = False
        self.wallpaper_interval = 60000
        self.show_clock = True
        self.wallpaper_dir = "wallpapers"
        os.makedirs(self.wallpaper_dir, exist_ok=True)
        self.cursor_folder = "cursors"
        os.makedirs(self.cursor_folder, exist_ok=True)
        self.download_wallpapers()
        self.wallpapers = [os.path.join(self.wallpaper_dir, f) for f in os.listdir(self.wallpaper_dir)]
        self.current_wallpaper = None
        self.cursor_path = ""
        self.apps = {}
        self.central_widget = QWidget(self)
        self.setCentralWidget(self.central_widget)
        self.central_widget.setAutoFillBackground(True)
        self.main_layout = QVBoxLayout(self.central_widget)
        self.main_layout.setContentsMargins(0, 0, 0, 0)
        self.main_layout.setSpacing(0)
        self.desktop_area = QWidget(self.central_widget)
        self.desktop_area.setStyleSheet("background: transparent;")
        self.desktop_layout = QGridLayout(self.desktop_area)
        self.desktop_layout.setSpacing(10)
        self.desktop_layout.setContentsMargins(10, 10, 10, 10)
        self.main_layout.addWidget(self.desktop_area)
        self.open_apps = []
        self.add_desktop_icons()
        self.taskbar = QWidget(self)
        self.taskbar.setStyleSheet("background-color: rgba(0,0,0,0.7);")
        self.taskbar.setFixedHeight(50)
        self.taskbar_layout = QHBoxLayout(self.taskbar)
        self.taskbar_layout.setContentsMargins(10, 5, 10, 5)
        self.taskbar_layout.setSpacing(10)
        self.start_button = QPushButton(self.start_text)
        self.start_button.setStyleSheet(
            "background-color: #333; border: 1px solid #555; border-radius: 10px; padding: 4px 6px; color: white; font-size: 16px;")
        self.start_button.setFixedSize(80, 40)
        self.start_button.clicked.connect(self.toggle_start_menu)
        self.taskbar_layout.addWidget(self.start_button, alignment=Qt.AlignmentFlag.AlignLeft)
        self.change_wallpaper_button = WallpaperButton(self.change_wallpaper_text)
        self.change_wallpaper_button.setStyleSheet(
            "background-color: #333; border: 1px solid #555; border-radius: 10px; padding: 4px 6px; color: white; font-size: 16px;")
        self.change_wallpaper_button.setFixedSize(120, 40)
        self.change_wallpaper_button.clicked.connect(self.change_wallpaper)
        self.change_wallpaper_button.onRightClick = self.open_wallpaper_chooser
        self.taskbar_layout.addWidget(self.change_wallpaper_button, alignment=Qt.AlignmentFlag.AlignLeft)
        self.app_button1 = QPushButton(self.file_explorer_text)
        self.app_button1.setStyleSheet(
            "background-color: #333; border: 1px solid #555; border-radius: 10px; padding: 4px 6px; color: white; font-size: 16px;")
        self.app_button1.setFixedSize(100, 40)
        # Запускаем новое окно проводника с дополнительными кнопками (DraggableFileExplorer)
        self.app_button1.clicked.connect(lambda: self.launch_app(DraggableFileExplorer, "file_explorer"))
        self.taskbar_layout.addWidget(self.app_button1, alignment=Qt.AlignmentFlag.AlignLeft)
        self.app_button2 = QPushButton(self.terminal_text)
        self.app_button2.setStyleSheet(
            "background-color: #333; border: 1px solid #555; border-radius: 10px; padding: 4px 6px; color: white; font-size: 16px;")
        self.app_button2.setFixedSize(100, 40)
        self.app_button2.clicked.connect(lambda: self.launch_app(self.open_terminal_app, "terminal"))
        self.taskbar_layout.addWidget(self.app_button2, alignment=Qt.AlignmentFlag.AlignLeft)
        self.app_button3 = QPushButton(self.settings_text)
        self.app_button3.setStyleSheet(
            "background-color: #333; border: 1px solid #555; border-radius: 10px; padding: 4px 6px; color: white; font-size: 16px;")
        self.app_button3.setFixedSize(100, 40)
        self.app_button3.clicked.connect(lambda: self.launch_app(self.open_settings, "settings"))
        self.taskbar_layout.addWidget(self.app_button3, alignment=Qt.AlignmentFlag.AlignLeft)
        self.cursor_button = QPushButton(self.cursor_text)
        self.cursor_button.setStyleSheet(
            "background-color: #333; border: 1px solid #555; border-radius: 10px; padding: 4px 6px; color: white; font-size: 16px;")
        self.cursor_button.setFixedSize(100, 40)
        self.cursor_button.clicked.connect(lambda: self.launch_app(self.open_cursor_app, "cursor"))
        self.taskbar_layout.addWidget(self.cursor_button, alignment=Qt.AlignmentFlag.AlignLeft)
        self.taskbar_layout.addStretch()
        self.battery_label = QLabel(f"{self.battery_text}: N/A")
        self.battery_label.setStyleSheet(
            "QLabel { background-color: #333; border: 1px solid #555; border-radius: 10px; padding: 2px 4px; color: white; font-size: 16px; }")
        self.taskbar_layout.addWidget(self.battery_label, alignment=Qt.AlignmentFlag.AlignRight)
        self.network_label = QLabel(f"{self.network_text}: N/A")
        self.network_label.setStyleSheet(
            "QLabel { background-color: #333; border: 1px solid #555; border-radius: 10px; padding: 2px 4px; color: white; font-size: 16px; }")
        self.taskbar_layout.addWidget(self.network_label, alignment=Qt.AlignmentFlag.AlignRight)
        self.time_button = QPushButton()
        self.time_button.setStyleSheet(
            "QPushButton { background-color: #333; border: 1px solid #555; border-radius: 10px; padding: 2px 4px; color: white; font-size: 16px; }")
        self.time_button.clicked.connect(self.show_calendar)
        self.taskbar_layout.addWidget(self.time_button, alignment=Qt.AlignmentFlag.AlignRight)
        self.main_layout.addWidget(self.taskbar, alignment=Qt.AlignmentFlag.AlignBottom)
        self.load_settings()
        if self.current_wallpaper and os.path.exists(self.current_wallpaper):
            pixmap = QPixmap(self.current_wallpaper)
            palette = self.central_widget.palette()
            palette.setBrush(QPalette.ColorRole.Window,
                             QBrush(pixmap.scaled(self.size(), Qt.AspectRatioMode.KeepAspectRatioByExpanding,
                                                  Qt.TransformationMode.SmoothTransformation)))
            self.central_widget.setPalette(palette)
        else:
            self.change_wallpaper()
        self.start_sound_sensor()
        self.update_clock()
        self.timer = QTimer(self)
        self.timer.timeout.connect(self.update_clock)
        self.timer.start(1000)
        if self.auto_wallpaper:
            self.wallpaper_timer = QTimer(self)
            self.wallpaper_timer.timeout.connect(self.change_wallpaper)
            self.wallpaper_timer.start(self.wallpaper_interval)
        self.shortcut_close = QShortcut(QKeySequence("Ctrl+Escape"), self)
        self.shortcut_close.activated.connect(self.close)
        self.start_menu = None
        self.show()

    def start_sound_sensor(self):
        if sd is None or np is None:
            return

        def callback(indata, frames, time, status):
            if status:
                print(status)
            rms = np.sqrt(np.mean(indata ** 2))
            level = min(100, int(rms * 5000))

        self.sound_stream = sd.InputStream(callback=callback)
        self.sound_stream.start()

    def load_settings(self):
        settings = load_json_settings()
        self.auto_wallpaper = settings.get("auto_wallpaper", False)
        self.wallpaper_interval = settings.get("wallpaper_interval", 60) * 1000
        self.show_clock = settings.get("show_clock", True)
        theme = settings.get("theme", "fusion")
        self.apply_theme(theme)
        saved_wp = settings.get("current_wallpaper", "")
        if saved_wp and os.path.exists(saved_wp):
            self.current_wallpaper = saved_wp
        self.cursor_path = settings.get("cursor", "")
        if self.cursor_path and os.path.exists(self.cursor_path):
            QApplication.setOverrideCursor(QCursor(QPixmap(self.cursor_path)))

    def apply_theme(self, theme_name):
        if theme_name.lower() == "fusion":
            QApplication.instance().setStyle("Fusion")
            QApplication.instance().setStyleSheet("")
        else:
            QApplication.instance().setStyle("")
            QApplication.instance().setStyleSheet("")

    def start_wallpaper_timer(self):
        if self.auto_wallpaper:
            if not hasattr(self, 'wallpaper_timer'):
                self.wallpaper_timer = QTimer(self)
                self.wallpaper_timer.timeout.connect(self.change_wallpaper)
            self.wallpaper_timer.start(self.wallpaper_interval)

    def stop_wallpaper_timer(self):
        if hasattr(self, "wallpaper_timer"):
            self.wallpaper_timer.stop()

    def add_wallpaper(self):
        file_path, _ = QFileDialog.getOpenFileName(self, "Выберите обои", "",
                                                   "Images (*.png *.jpg *.jpeg *.bmp *.gif)")
        if file_path:
            base_name = os.path.basename(file_path)
            dest_path = os.path.join(self.wallpaper_dir,
                                     f"custom_{datetime.now().strftime('%Y%m%d%H%M%S')}_{base_name}")
            try:
                with open(file_path, "rb") as src, open(dest_path, "wb") as dst:
                    dst.write(src.read())
                QMessageBox.information(self, "Обои", "Обои успешно добавлены!")
                self.wallpapers.append(dest_path)
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", f"Не удалось добавить обои: {str(e)}")

    def add_desktop_icons(self):
        apps = [
            (self.file_explorer_text, DraggableFileExplorer, "file_explorer"),
            (self.terminal_text, self.open_terminal_app, "terminal"),
            (self.settings_text, self.open_settings, "settings"),
            ("Браузер", App4, "Браузер"),
            ("Калькулятор", App5, "Калькулятор"),
        ]
        cols = 5
        for index, (name, func, key) in enumerate(apps):
            if index >= 20:
                break
            btn = QPushButton(name)
            btn.setStyleSheet(
                "background-color: #333; border: 1px solid #555; border-radius: 10px; padding: 4px 6px; color: white; font-size: 16px;")
            btn.setFixedSize(100, 60)
            btn.clicked.connect(lambda _, f=func, k=key: self.launch_app(f, k))
            row = index // cols
            col = index % cols
            self.desktop_layout.addWidget(btn, row, col, alignment=Qt.AlignmentFlag.AlignCenter)

    def launch_app(self, app_function, key):
        if key in self.apps and self.apps[key]:
            instance = self.apps[key]
            if instance.isMinimized():
                instance.showNormal()
                instance.raise_()
                instance.activateWindow()
            else:
                instance.showMinimized()
            return
        inst = app_function()
        inst.show()
        self.apps[key] = inst
        inst.destroyed.connect(lambda: self.apps.pop(key, None))

    def toggle_start_menu(self):
        if self.start_menu and self.start_menu.isVisible():
            self.start_menu.close()
        else:
            self.show_start_menu()

    def show_start_menu(self):
        self.start_menu = QDialog(self)
        self.start_menu.setWindowFlags(Qt.WindowType.FramelessWindowHint)
        self.start_menu.setStyleSheet("background-color: rgba(0,0,0,0.9); border: 2px solid white;")
        self.start_menu.resize(300, 400)
        screen_geom = QApplication.primaryScreen().availableGeometry()
        self.start_menu.move(screen_geom.width() // 2 - 150, screen_geom.height() // 2 - 200)
        layout = QVBoxLayout(self.start_menu)
        layout.setContentsMargins(10, 10, 10, 10)
        layout.setSpacing(10)
        buttons = {
            self.file_explorer_text: ("file_explorer", DraggableFileExplorer),
            self.terminal_text: ("terminal", self.open_terminal_app),
            self.settings_text: ("settings", self.open_settings),
            self.cursor_text: ("cursor", self.open_cursor_app),
            "Выход": ("exit", self.close)
        }
        for name, (key, act) in buttons.items():
            btn = QPushButton(name)
            btn.setStyleSheet(
                "background-color: #333; border: 1px solid #555; border-radius: 10px; padding: 4px 6px; color: white; font-size: 16px;")
            btn.setFixedSize(120, 50)
            btn.clicked.connect(lambda _, act=act, k=key: act() if k == "exit" else self.launch_app(act, k))
            layout.addWidget(btn, alignment=Qt.AlignmentFlag.AlignCenter)
        self.start_menu.show()

    def open_file_explorer(self):
        return DraggableFileExplorer()

    def open_terminal_app(self):
        term = TerminalApp(desktop_window=self)
        return term

    def open_settings(self, developer=False):
        sett = SettingsWindow(self, developer=developer)
        return sett

    def open_hidden_app(self):
        return HiddenApp()

    def open_cursor_app(self):
        app = CursorChooserApp(parent=self)
        return app

    def change_wallpaper(self):
        if not self.wallpapers:
            return
        new_wp = choice(self.wallpapers)
        while new_wp == self.current_wallpaper and len(self.wallpapers) > 1:
            new_wp = choice(self.wallpapers)
        self.current_wallpaper = new_wp
        pixmap = QPixmap(self.current_wallpaper)
        palette = self.central_widget.palette()
        palette.setBrush(QPalette.ColorRole.Window, QBrush(pixmap.scaled(self.size(),
                                                                         Qt.AspectRatioMode.KeepAspectRatioByExpanding,
                                                                         Qt.TransformationMode.SmoothTransformation)))
        self.central_widget.setPalette(palette)
        self.save_current_wallpaper()

    def open_wallpaper_chooser(self):
        chooser = WallpaperChooserDialog(self.wallpapers, parent=self)
        if chooser.exec() == QDialog.DialogCode.Accepted and chooser.selected_wallpaper:
            self.current_wallpaper = chooser.selected_wallpaper
            pixmap = QPixmap(self.current_wallpaper)
            palette = self.central_widget.palette()
            palette.setBrush(QPalette.ColorRole.Window, QBrush(pixmap.scaled(self.size(),
                                                                             Qt.AspectRatioMode.KeepAspectRatioByExpanding,
                                                                             Qt.TransformationMode.SmoothTransformation)))
            self.central_widget.setPalette(palette)
            self.save_current_wallpaper()

    def save_current_wallpaper(self):
        settings = load_json_settings()
        settings["current_wallpaper"] = self.current_wallpaper if self.current_wallpaper else ""
        save_json_settings(settings)

    def save_current_cursor(self):
        from __main__ import save_current_cursor
        save_current_cursor(self)

    def download_wallpapers(self):
        wallpaper_urls = [
            "https://7themes.su/_ph/65/467656588.jpg",
            "https://www.fonstola.ru/images/202204/www.fonstola.ru.1649743041.3966.jpg",
            "https://img.goodfon.ru/original/3840x2160/6/77/windows-11-vaporwave-desktop-wallpaper-hd-purple.jpg",
            "https://images.wallpaperscraft.ru/image/single/mashina_neon_podsvetka_158672_3840x2160.jpg",
            "https://cs9.pikabu.ru/post_img/big/2020/05/12/9/1589296652173879650.jpg",
            "https://cs4.pikabu.ru/post_img/big/2016/05/26/6/1464254142114738279.jpg",
            "https://cs4.pikabu.ru/post_img/big/2016/05/26/6/1464254142114738279.jpg",
            "https://image.fonwall.ru/o/ey/wallpaper-anime-landscape-cityscape-scenic-sunset.jpeg",
            "https://images.wallpaperscraft.ru/image/single/ulitsa_noch_mokryj_155637_3840x2160.jpg",
        ]
        headers = {"User-Agent": "Mozilla/5.0"}
        for i, url in enumerate(wallpaper_urls):
            local_path = os.path.join(self.wallpaper_dir, f"wallpaper_{i + 1}.jpg")
            if not os.path.exists(local_path):
                try:
                    response = requests.get(url, headers=headers)
                    response.raise_for_status()
                    image = Image.open(BytesIO(response.content))
                    image = image.resize((1920, 1080), Image.Resampling.LANCZOS)
                    image.save(local_path)
                except Exception as e:
                    print("Не удалось скачать обои: " + url + "\nОшибка: " + str(e))

    def update_clock(self):
        current_time = QTime.currentTime().toString("HH:mm:ss")
        self.time_button.setText(current_time)
        battery = psutil.sensors_battery()
        if battery:
            battery_status = f"{self.battery_text}: {battery.percent}%"
            if battery.power_plugged:
                battery_status += " (Заряжается)"
            else:
                if battery.secsleft not in (psutil.POWER_TIME_UNKNOWN, psutil.POWER_TIME_UNLIMITED):
                    secs = battery.secsleft
                    hrs, rem = divmod(secs, 3600)
                    mins, secs = divmod(rem, 60)
                    battery_status += f" (Осталось примерно {hrs} ч {mins} м)"
                else:
                    battery_status += " (Время разряда недоступно)"
        else:
            battery_status = f"{self.battery_text}: N/A"
        self.battery_label.setText(battery_status)
        net_if = psutil.net_if_stats()
        connected = any(stats.isup for iface, stats in net_if.items() if iface.lower() != "lo")
        network_status = f"{self.network_text}: Подключена" if connected else f"{self.network_text}: Отключена"
        self.network_label.setText(network_status)

    def show_calendar(self):
        self.calendar_dialog = CalendarWindow(self)
        screen_geom = QApplication.primaryScreen().availableGeometry()
        dialog_width, dialog_height = 400, 400
        center_x = (screen_geom.width() - dialog_width) // 2
        center_y = (screen_geom.height() - dialog_height) // 2
        self.calendar_dialog.setFixedSize(dialog_width, dialog_height)
        self.calendar_dialog.move(center_x, center_y)
        self.calendar_dialog.show()

    def closeEvent(self, event):
        self.save_current_wallpaper()
        super().closeEvent(event)
        QApplication.quit()

    def keyPressEvent(self, event):
        if event.key() == Qt.Key.Key_Escape and QApplication.keyboardModifiers() == Qt.KeyboardModifier.ControlModifier:
            self.close()
        else:
            super().keyPressEvent(event)


class ImageEditor(QDialog):
    def __init__(self, file_path, parent=None):
        super().__init__(parent)
        self.setWindowTitle("Редактор изображения")
        self.file_path = file_path
        self.image = QImage(file_path)
        if self.image.isNull():
            QMessageBox.critical(self, "Ошибка", "Не удалось загрузить изображение.")
            self.close()
            return

        self.resize(self.image.size())
        self.label = QLabel()
        self.label.setAlignment(Qt.AlignmentFlag.AlignCenter)
        self.label.setPixmap(QPixmap.fromImage(self.image))
        scroll_area = QScrollArea()
        scroll_area.setWidget(self.label)
        scroll_area.setWidgetResizable(True)

        btn_rotate = QPushButton("Повернуть 90°")
        btn_rotate.clicked.connect(self.rotate_image)
        btn_flip = QPushButton("Отразить по горизонтали")
        btn_flip.clicked.connect(self.flip_image)
        btn_save = QPushButton("Сохранить")
        btn_save.clicked.connect(self.save_image)
        btn_close = QPushButton("Закрыть")
        btn_close.clicked.connect(self.close)

        btn_layout = QHBoxLayout()
        btn_layout.addWidget(btn_rotate)
        btn_layout.addWidget(btn_flip)
        btn_layout.addWidget(btn_save)
        btn_layout.addWidget(btn_close)

        layout = QVBoxLayout(self)
        layout.addWidget(scroll_area)
        layout.addLayout(btn_layout)

    def rotate_image(self):
        transform = QTransform().rotate(90)
        self.image = self.image.transformed(transform)
        self.label.setPixmap(QPixmap.fromImage(self.image))

    def flip_image(self):
        self.image = self.image.mirrored(True, False)
        self.label.setPixmap(QPixmap.fromImage(self.image))

    def save_image(self):
        save_path, _ = QFileDialog.getSaveFileName(self, "Сохранить изображение", self.file_path,
                                                   "Images (*.png *.jpg *.bmp)")
        if save_path:
            if self.image.save(save_path):
                QMessageBox.information(self, "Сохранено", "Изображение успешно сохранено.")
            else:
                QMessageBox.critical(self, "Ошибка", "Не удалось сохранить изображение.")


class FileTable(QTableWidget):
    def __init__(self, parent=None):
        super().__init__(parent)
        self.setSelectionMode(QAbstractItemView.SelectionMode.ExtendedSelection)
        self.setDragEnabled(True)
        self.setAcceptDrops(True)
        self.setDropIndicatorShown(True)

    def dragEnterEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
        else:
            super().dragEnterEvent(event)

    def dragMoveEvent(self, event):
        if event.mimeData().hasUrls():
            event.acceptProposedAction()
        else:
            super().dragMoveEvent(event)

    def dropEvent(self, event):
        if event.mimeData().hasUrls():
            urls = event.mimeData().urls()
            target_dir = self.parent().current_path
            for url in urls:
                source = url.toLocalFile()
                if source:
                    try:
                        if os.path.isdir(source):
                            shutil.copytree(source, os.path.join(target_dir, os.path.basename(source)))
                        else:
                            shutil.copy2(source, target_dir)
                    except Exception as e:
                        QMessageBox.critical(self, "Ошибка", f"Не удалось скопировать {source}: {str(e)}")
            event.acceptProposedAction()
        else:
            super().dropEvent(event)


class FileViewer(QDialog):
    def __init__(self, file_path, parent=None):
        super().__init__(parent)
        self.setModal(False)
        self.file_path = file_path
        self.setWindowTitle(os.path.basename(file_path))
        self.resize(800, 600)
        layout = QVBoxLayout(self)
        ext = QFileInfo(file_path).suffix().lower()

        image_extensions = ['png', 'jpg', 'jpeg', 'bmp', 'gif', 'ico']
        editable_extensions = ['txt', 'py', 'log', 'md', 'json', 'xml', 'html', 'csv', 'ini', 'cfg', 'cpp', 'c', 'java',
                               'js', 'css', 'php']
        zip_extensions = ['zip']
        db_extensions = ['db', 'sqlite']
        music_extensions = ['mp3', 'wav', 'ogg', 'flac']
        video_extensions = ['mp4', 'avi', 'mkv', 'mov']
        office_extensions = ['docx', 'doxc', 'pptx', 'xlsx']

        if ext in image_extensions:
            pixmap = QPixmap(file_path)
            self.resize(pixmap.size())
            scroll_area = QScrollArea()
            self.image_label = QLabel()
            self.image_label.setAlignment(Qt.AlignmentFlag.AlignCenter)
            self.image_label.setPixmap(pixmap)
            scroll_area.setWidget(self.image_label)
            scroll_area.setWidgetResizable(True)
            layout.addWidget(scroll_area)
            btn_editor = QPushButton("Открыть редактор")
            btn_editor.clicked.connect(lambda: self.open_image_editor(file_path))
            layout.addWidget(btn_editor)
        elif ext in video_extensions:
            self.media_player = QMediaPlayer(self)
            video_widget = QVideoWidget()
            self.media_player.setVideoOutput(video_widget)
            self.media_player.setSource(QUrl.fromLocalFile(file_path))
            layout.addWidget(video_widget)
            control_layout = QHBoxLayout()
            btn_play = QPushButton("Play")
            btn_pause = QPushButton("Pause")
            btn_stop = QPushButton("Stop")
            btn_play.clicked.connect(self.media_player.play)
            btn_pause.clicked.connect(self.media_player.pause)
            btn_stop.clicked.connect(self.media_player.stop)
            control_layout.addWidget(btn_play)
            control_layout.addWidget(btn_pause)
            control_layout.addWidget(btn_stop)
            layout.addLayout(control_layout)
        elif ext == 'pdf' and pdf_supported:
            pdf_doc = QPdfDocument(self)
            pdf_doc.load(file_path)
            pdf_view = QPdfView(self)
            pdf_view.setDocument(pdf_doc)
            layout.addWidget(pdf_view)
        elif ext in editable_extensions:
            if ext == 'html':
                text_edit = QTextEdit()
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                    text_edit.setText(content)
                except Exception as e:
                    text_edit.setText(f"Ошибка при открытии файла: {str(e)}")
                text_edit.setReadOnly(True)
                layout.addWidget(text_edit)
            else:
                self.text_edit = QTextEdit()
                try:
                    with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                        content = f.read()
                    self.text_edit.setText(content)
                except Exception as e:
                    self.text_edit.setText(f"Ошибка при открытии файла: {str(e)}")
                self.text_edit.setReadOnly(not os.access(file_path, os.W_OK))
                layout.addWidget(self.text_edit)
                btn_layout = QHBoxLayout()
                self.btn_save = QPushButton("Сохранить")
                self.btn_save.clicked.connect(self.save_file)
                btn_close = QPushButton("Закрыть")
                btn_close.clicked.connect(self.close)
                btn_layout.addWidget(self.btn_save)
                btn_layout.addWidget(btn_close)
                layout.addLayout(btn_layout)
        elif ext in zip_extensions:
            self.zip_list = QListWidget()
            try:
                with zipfile.ZipFile(file_path, 'r') as zip_ref:
                    self.zip_file_list = zip_ref.namelist()
                self.zip_list.addItems(self.zip_file_list)
                self.zip_list.itemDoubleClicked.connect(self.open_zip_item)
                layout.addWidget(self.zip_list)
            except Exception as e:
                layout.addWidget(QLabel(f"Ошибка при открытии ZIP файла: {str(e)}"))
        elif ext in db_extensions:
            try:
                conn = sqlite3.connect(file_path)
                cursor = conn.cursor()
                cursor.execute("SELECT name FROM sqlite_master WHERE type='table';")
                tables = cursor.fetchall()
                conn.close()
                text_edit = QTextEdit()
                text_edit.setReadOnly(True)
                if tables:
                    table_list = "\n".join([t[0] for t in tables])
                    text_edit.setText("Таблицы базы данных:\n" + table_list)
                else:
                    text_edit.setText("База данных пуста или не содержит таблиц.")
                layout.addWidget(text_edit)
            except Exception as e:
                layout.addWidget(QLabel(f"Ошибка при чтении базы данных: {str(e)}"))
        elif ext in music_extensions:
            self.media_player = QMediaPlayer(self)
            self.audio_output = QAudioOutput(self)
            self.media_player.setAudioOutput(self.audio_output)
            self.media_player.setSource(QUrl.fromLocalFile(file_path))
            control_layout = QHBoxLayout()
            btn_play = QPushButton("Play")
            btn_pause = QPushButton("Pause")
            btn_stop = QPushButton("Stop")
            btn_play.clicked.connect(self.media_player.play)
            btn_pause.clicked.connect(self.media_player.pause)
            btn_stop.clicked.connect(self.media_player.stop)
            control_layout.addWidget(btn_play)
            control_layout.addWidget(btn_pause)
            control_layout.addWidget(btn_stop)
            layout.addLayout(control_layout)
        elif ext in office_extensions:
            try:
                with zipfile.ZipFile(file_path, 'r') as zip_ref:
                    if ext in ['docx', 'doxc']:
                        with zip_ref.open('word/document.xml') as doc_xml:
                            xml_content = doc_xml.read()
                        root = ET.fromstring(xml_content)
                        ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
                        texts = [node.text for node in root.findall('.//w:t', ns) if node.text]
                        full_text = "\n".join(texts)
                    elif ext == 'pptx':
                        full_text = ""
                        slide_names = [name for name in zip_ref.namelist() if name.startswith("ppt/slides/slide")]
                        for slide_name in slide_names:
                            with zip_ref.open(slide_name) as slide_xml:
                                xml_content = slide_xml.read()
                            root = ET.fromstring(xml_content)
                            ns = {'a': 'http://schemas.openxmlformats.org/drawingml/2006/main',
                                  'p': 'http://schemas.openxmlformats.org/presentationml/2006/main'}
                            texts = [node.text for node in root.findall('.//a:t', ns) if node.text]
                            full_text += "\n".join(texts) + "\n"
                    elif ext == 'xlsx':
                        full_text = "Предварительный просмотр Excel файлов пока не поддерживается."
                    else:
                        full_text = "Формат файла не поддерживается для предварительного просмотра."
                text_edit = QTextEdit()
                text_edit.setReadOnly(True)
                text_edit.setText(full_text)
                layout.addWidget(text_edit)
            except Exception as e:
                layout.addWidget(QLabel(f"Ошибка при чтении Office документа: {str(e)}"))
        else:
            try:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    content = f.read()
                text_edit = QTextEdit()
                text_edit.setReadOnly(True)
                text_edit.setText(content)
                layout.addWidget(text_edit)
            except Exception as e:
                layout.addWidget(QLabel("Формат файла не поддерживается для просмотра."))

    def open_zip_item(self, item):
        file_name = item.text()
        try:
            with zipfile.ZipFile(self.file_path, 'r') as zip_ref:
                data = zip_ref.read(file_name)
            suffix = os.path.splitext(file_name)[1]
            tmp = tempfile.NamedTemporaryFile(delete=False, suffix=suffix)
            tmp.write(data)
            tmp.close()
            viewer = FileViewer(tmp.name, self)
            viewer.show()
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось открыть файл из архива: {str(e)}")

    def open_image_editor(self, file_path):
        editor = ImageEditor(file_path, self)
        editor.exec()

    def save_file(self):
        try:
            with open(self.file_path, 'w', encoding='utf-8') as f:
                f.write(self.text_edit.toPlainText())
            QMessageBox.information(self, "Сохранено", "Изменения успешно сохранены.")
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Не удалось сохранить файл: {str(e)}")

    def closeEvent(self, event):
        if hasattr(self, 'media_player'):
            self.media_player.stop()
        event.accept()


class FileExplorer(QMainWindow):
    def __init__(self):
        super().__init__()
        self.settings = QSettings("FileManager", "Explorer")
        self.setWindowTitle("Проводник")
        self.setMinimumSize(1024, 768)
        self.history = []
        self.history_index = -1
        self.current_path = self.settings.value("last_path", QDir.homePath())
        self.entries = []
        self.animated_index = 0
        self.animation_timer = QTimer(self)
        self.animation_timer.timeout.connect(self.add_next_row_animated)
        self.init_ui()
        self.load_directory(self.current_path, animate=True)

    def init_ui(self):
        main_widget = QWidget()
        self.setCentralWidget(main_widget)
        self.layout = QVBoxLayout(main_widget)
        self.create_toolbar()
        self.create_file_table()
        self.statusBar().showMessage("Готово")

    def create_toolbar(self):
        toolbar = QHBoxLayout()
        self.btn_back = QToolButton()
        self.btn_back.setIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_ArrowBack))
        self.btn_back.clicked.connect(self.navigate_back)
        self.btn_forward = QToolButton()
        self.btn_forward.setIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_ArrowForward))
        self.btn_forward.clicked.connect(self.navigate_forward)
        self.btn_up = QToolButton()
        self.btn_up.setText("Вверх")
        self.btn_up.clicked.connect(self.navigate_up)
        self.path_edit = QLineEdit()
        self.path_edit.setText(self.current_path)
        self.path_edit.returnPressed.connect(partial(self.load_directory, self.path_edit.text()))
        self.btn_refresh = QToolButton()
        self.btn_refresh.setIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_BrowserReload))
        self.btn_refresh.clicked.connect(lambda: self.load_directory(self.current_path, animate=True))
        self.btn_new = QToolButton()
        self.btn_new.setIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_FileDialogNewFolder))
        self.btn_new.setMenu(self.create_new_menu())
        self.btn_new.setPopupMode(QToolButton.ToolButtonPopupMode.InstantPopup)

        self.btn_devices = QToolButton()
        self.btn_devices.setIcon(self.style().standardIcon(QStyle.StandardPixmap.SP_DriveFDIcon))
        self.btn_devices.setText("Устройства")
        self.btn_devices.setPopupMode(QToolButton.ToolButtonPopupMode.InstantPopup)
        devices_menu = QMenu()
        self.btn_devices.setMenu(devices_menu)
        devices_menu.aboutToShow.connect(self.update_devices_menu)

        toolbar.addWidget(self.btn_back)
        toolbar.addWidget(self.btn_forward)
        toolbar.addWidget(self.btn_up)
        toolbar.addWidget(self.path_edit)
        toolbar.addWidget(self.btn_refresh)
        toolbar.addWidget(self.btn_new)
        toolbar.addWidget(self.btn_devices)
        self.layout.addLayout(toolbar)

    def create_new_menu(self):
        menu = QMenu()
        actions = [
            ("Папку", self.create_folder),
            ("Файл", self.create_file)
        ]
        for text, callback in actions:
            action = QAction(text, self)
            action.triggered.connect(callback)
            menu.addAction(action)
        return menu

    def update_devices_menu(self):
        menu = self.btn_devices.menu()
        menu.clear()
        drives = QStorageInfo.mountedVolumes()
        available_drives = []
        for drive in drives:
            if drive.isValid() and drive.isReady():
                # Формируем имя тома: если есть отображаемое имя, используем его, иначе — путь
                root = drive.rootPath()
                name = drive.displayName() if drive.displayName() else root
                available_drives.append((name, root))
        if available_drives:
            # Если есть хотя бы один том, добавляем их в меню
            for name, root in available_drives:
                action = QAction(name, self)
                # Используем замыкание для передачи параметра root в лямбду
                action.triggered.connect(lambda checked, p=root: self.load_directory(p, animate=True))
                menu.addAction(action)
        else:
            # Если нет доступных томов, добавляем действие для перехода к системному диску
            current_drive = os.path.splitdrive(self.current_path)[0] + os.sep
            action = QAction("Локальный диск (" + current_drive + ")", self)
            action.triggered.connect(lambda: self.load_directory(current_drive, animate=True))
            menu.addAction(action)

    def create_file_table(self):
        self.table = FileTable(self)
        self.table.setColumnCount(4)
        self.table.setHorizontalHeaderLabels(["Имя", "Тип", "Размер", "Изменен"])
        self.table.setSortingEnabled(True)
        self.table.horizontalHeader().setSectionResizeMode(0, QHeaderView.ResizeMode.Stretch)
        self.table.doubleClicked.connect(self.open_item)
        self.table.setContextMenuPolicy(Qt.ContextMenuPolicy.CustomContextMenu)
        self.table.customContextMenuRequested.connect(self.show_context_menu)
        self.layout.addWidget(self.table)

    def show_context_menu(self, pos):
        menu = QMenu()
        actions = [
            ("Открыть", self.open_selected),
            ("Удалить", self.delete_selected),
            ("Копировать", self.copy_selected),
            ("Переместить", self.move_selected),
            ("Переименовать", self.rename_selected),
            ("Свойства", self.show_properties)
        ]
        for text, callback in actions:
            action = QAction(text, self)
            action.triggered.connect(callback)
            menu.addAction(action)
        menu.exec(self.table.viewport().mapToGlobal(pos))

    def load_directory(self, path, animate=True):
        try:
            if not QFileInfo(path).exists():
                raise FileNotFoundError
            self.update_history(path)
            self.current_path = path
            self.path_edit.setText(path)
            self.settings.setValue("last_path", path)
            self.update_navigation_buttons()
            entries = QDir(path).entryInfoList(
                QDir.Filter.AllEntries | QDir.Filter.NoDotAndDotDot,
                QDir.SortFlag.DirsFirst | QDir.SortFlag.IgnoreCase
            )
            if not animate:
                entries = sorted(entries, key=lambda info: info.lastModified(), reverse=True)
                self.table.setRowCount(len(entries))
                for row, info in enumerate(entries):
                    self.add_table_row(row, info)
            else:
                self.entries = entries
                self.animated_index = 0
                self.table.setRowCount(0)
                self.animation_timer.start(5)
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", f"Ошибка доступа: {str(e)}")

    def add_next_row_animated(self):
        if self.animated_index < len(self.entries):
            info = self.entries[self.animated_index]
            row = self.table.rowCount()
            self.table.insertRow(row)
            self.add_table_row(row, info)
            self.animated_index += 1
        else:
            self.animation_timer.stop()

    def add_table_row(self, row, file_info):
        icon = self.style().standardIcon(
            QStyle.StandardPixmap.SP_DirIcon if file_info.isDir() else QStyle.StandardPixmap.SP_FileIcon
        )
        name_item = QTableWidgetItem(icon, file_info.fileName())
        type_item = QTableWidgetItem("Папка" if file_info.isDir() else file_info.suffix().upper())
        size_item = QTableWidgetItem(self.format_size(file_info.size()) if file_info.isFile() else "")
        date_item = QTableWidgetItem(file_info.lastModified().toString("dd.MM.yyyy HH:mm"))
        self.table.setItem(row, 0, name_item)
        self.table.setItem(row, 1, type_item)
        self.table.setItem(row, 2, size_item)
        self.table.setItem(row, 3, date_item)
        name_item.setData(Qt.ItemDataRole.UserRole, file_info.absoluteFilePath())

    def format_size(self, size):
        units = ["Б", "КБ", "МБ", "ГБ"]
        for unit in units:
            if size < 1024:
                return f"{size:.1f} {unit}"
            size /= 1024
        return f"{size:.1f} ТБ"

    def update_history(self, path):
        if self.history and self.history[self.history_index] == path:
            return
        self.history = self.history[:self.history_index + 1]
        self.history.append(path)
        self.history_index += 1

    def navigate_back(self):
        if self.history_index > 0:
            self.history_index -= 1
            self.load_directory(self.history[self.history_index], animate=True)

    def navigate_forward(self):
        if self.history_index < len(self.history) - 1:
            self.history_index += 1
            self.load_directory(self.history[self.history_index], animate=True)

    def navigate_up(self):
        parent = os.path.dirname(self.current_path)
        if parent and parent != self.current_path:
            self.load_directory(parent, animate=True)

    def update_navigation_buttons(self):
        self.btn_back.setEnabled(self.history_index > 0)
        self.btn_forward.setEnabled(self.history_index < len(self.history) - 1)

    def create_folder(self):
        name, ok = QInputDialog.getText(self, "Новая папка", "Имя папки:")
        if ok and name:
            try:
                os.mkdir(os.path.join(self.current_path, name))
                self.load_directory(self.current_path, animate=False)
                QMessageBox.information(self, "Создано", f"Создана папка: {name}")
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", str(e))

    def create_file(self):
        name, ok = QInputDialog.getText(self, "Новый файл", "Имя файла:")
        if ok and name:
            try:
                open(os.path.join(self.current_path, name), 'a').close()
                self.load_directory(self.current_path, animate=False)
                QMessageBox.information(self, "Создано", f"Создан файл: {name}")
            except Exception as e:
                QMessageBox.critical(self, "Ошибка", str(e))

    def delete_selected(self):
        selected = self.table.selectedItems()
        if selected:
            paths = list({item.data(Qt.ItemDataRole.UserRole) for item in selected})
            reply = QMessageBox.question(
                self, "Удаление",
                f"Удалить выбранные файлы/папки?",
                QMessageBox.StandardButton.Yes | QMessageBox.StandardButton.No
            )
            if reply == QMessageBox.StandardButton.Yes:
                for path in paths:
                    self.perform_deletion(path)

    def copy_selected(self):
        selected = self.table.selectedItems()
        if selected:
            paths = list({item.data(Qt.ItemDataRole.UserRole) for item in selected})
            target = QFileDialog.getExistingDirectory(self, "Выберите папку для копирования", self.current_path)
            if target:
                for path in paths:
                    try:
                        if os.path.isdir(path):
                            shutil.copytree(path, os.path.join(target, os.path.basename(path)))
                        else:
                            shutil.copy2(path, target)
                    except Exception as e:
                        QMessageBox.critical(self, "Ошибка", f"Ошибка копирования {path}: {str(e)}")
                self.load_directory(self.current_path, animate=True)

    def move_selected(self):
        selected = self.table.selectedItems()
        if selected:
            paths = list({item.data(Qt.ItemDataRole.UserRole) for item in selected})
            target = QFileDialog.getExistingDirectory(self, "Выберите папку для перемещения", self.current_path)
            if target:
                for path in paths:
                    try:
                        shutil.move(path, target)
                    except Exception as e:
                        QMessageBox.critical(self, "Ошибка", f"Ошибка перемещения {path}: {str(e)}")
                self.load_directory(self.current_path, animate=True)

    def perform_deletion(self, path):
        progress = QProgressDialog("Удаление...", "Отмена", 0, 0, self)
        progress.setWindowModality(Qt.WindowModality.WindowModal)
        progress.show()
        try:
            if os.path.isdir(path):
                os.rmdir(path)
            else:
                os.remove(path)
        except Exception as e:
            QMessageBox.critical(self, "Ошибка", str(e))
        finally:
            progress.close()
            self.load_directory(self.current_path, animate=True)

    def open_item(self, index):
        path = self.table.item(index.row(), 0).data(Qt.ItemDataRole.UserRole)
        if QFileInfo(path).isDir():
            self.load_directory(path, animate=True)
        else:
            self.open_file(path)

    def open_file(self, path):
        viewer = FileViewer(path, self)
        viewer.show()

    def open_selected(self):
        selected = self.table.selectedItems()
        if selected:
            paths = list({item.data(Qt.ItemDataRole.UserRole) for item in selected})
            if len(paths) == 1:
                if QFileInfo(paths[0]).isDir():
                    self.load_directory(paths[0], animate=True)
                else:
                    self.open_file(paths[0])
            else:
                QMessageBox.information(self, "Информация", "Открытие нескольких файлов не поддерживается.")

    def rename_selected(self):
        selected = self.table.selectedItems()
        if selected:
            old_path = selected[0].data(Qt.ItemDataRole.UserRole)
            new_name, ok = QInputDialog.getText(self, "Переименовать", "Новое имя:")
            if ok and new_name:
                try:
                    new_path = os.path.join(os.path.dirname(old_path), new_name)
                    os.rename(old_path, new_path)
                    self.load_directory(self.current_path, animate=True)
                    QMessageBox.information(self, "Переименовано", f"Файл переименован в: {new_name}")
                except Exception as e:
                    QMessageBox.critical(self, "Ошибка", str(e))

    def show_properties(self):
        selected = self.table.selectedItems()
        if selected:
            path = selected[0].data(Qt.ItemDataRole.UserRole)
            info = QFileInfo(path)
            msg = QMessageBox()
            msg.setWindowTitle("Свойства")
            msg.setText(f"""
Имя: {info.fileName()}
Тип: {'Папка' if info.isDir() else 'Файл'}
Размер: {self.format_size(info.size())}
Дата изменения: {info.lastModified().toString("dd.MM.yyyy HH:mm")}
            """)
            msg.exec()


class DraggableFileExplorer(DraggableDialog):
    def __init__(self):
        super().__init__("Проводник")
        self.resize(1024, 768)
        self.explorer = FileExplorer()
        self.explorer.setWindowFlags(Qt.WindowType.Widget)

        layout = QVBoxLayout()
        layout.addWidget(self.explorer)
        self.setContentLayout(layout)


if __name__ == "__main__":
    app = QApplication(sys.argv)
    app.setStyle("Fusion")
    window = DesktopWindow()
    window.show()
    sys.exit(app.exec())
